#!/usr/bin/env python3
"""Generate and validate technical materials for China software copyright registration.

This tool intentionally does not create legal-entity documents or infer legal facts.
It is designed to be invoked by an agent following the bundled SKILL.md.
"""

from __future__ import annotations

import argparse
import csv
import datetime as dt
import fnmatch
import hashlib
import json
import os
import plistlib
import re
import shutil
import subprocess
import sys
import tempfile
import textwrap
import zipfile
from dataclasses import dataclass, field
from pathlib import Path, PurePosixPath
from typing import Any, Iterable, Sequence

try:
    import tomllib
except ModuleNotFoundError:  # Python < 3.11
    import tomli as tomllib

try:
    from docx import Document
    from docx.enum.section import WD_SECTION
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Inches, Pt, RGBColor
except ImportError:  # pragma: no cover
    Document = None  # type: ignore[assignment]

try:
    from pypdf import PdfReader, PdfWriter
except ImportError:  # pragma: no cover
    PdfReader = PdfWriter = None  # type: ignore[assignment]

try:
    from PIL import Image
except ImportError:  # pragma: no cover
    Image = None  # type: ignore[assignment]


RULES_CHECKED = dt.date(2026, 6, 22)
TOOL_VERSION = "2.1.0"
DEFAULT_EXTENSIONS = {
    ".java", ".kt", ".kts", ".swift", ".m", ".mm", ".dart", ".py", ".go",
    ".rs", ".c", ".h", ".cpp", ".hpp", ".cs", ".js", ".jsx", ".ts",
    ".tsx", ".vue", ".php", ".rb", ".scala", ".sql", ".xml", ".html",
    ".css", ".scss", ".less", ".sh", ".bash", ".lua", ".r", ".fs", ".fsx",
}
LANGUAGE_BY_EXT = {
    ".java": "Java", ".kt": "Kotlin", ".kts": "Kotlin", ".swift": "Swift",
    ".m": "Objective-C", ".mm": "Objective-C++", ".dart": "Dart", ".py": "Python",
    ".go": "Go", ".rs": "Rust", ".c": "C", ".h": "C/C++ Header", ".cpp": "C++",
    ".hpp": "C++ Header", ".cs": "C#", ".js": "JavaScript", ".jsx": "JavaScript/JSX",
    ".ts": "TypeScript", ".tsx": "TypeScript/TSX", ".vue": "Vue", ".php": "PHP",
    ".rb": "Ruby", ".scala": "Scala", ".sql": "SQL", ".xml": "XML",
    ".html": "HTML", ".css": "CSS", ".scss": "SCSS", ".less": "Less",
    ".sh": "Shell", ".bash": "Shell", ".lua": "Lua", ".r": "R", ".fs": "F#",
    ".fsx": "F#",
}
SENSITIVE_SUFFIXES = {".pem", ".key", ".p12", ".pfx", ".jks", ".keystore", ".crt", ".cer"}
PLACEHOLDER_RE = re.compile(r"(?i)\bTODO\b|\bTBD\b|待补充|占位|示意图占位|\{\{[^}]+\}\}")
HIGH_SECRET_PATTERNS = [
    ("private-key", re.compile(r"-----BEGIN (?:RSA |EC |OPENSSH |DSA )?PRIVATE KEY-----")),
    ("aws-access-key", re.compile(r"\bAKIA[0-9A-Z]{16}\b")),
    ("github-token", re.compile(r"\bgh[pousr]_[A-Za-z0-9_]{30,}\b")),
    ("google-api-key", re.compile(r"\bAIza[0-9A-Za-z\-_]{35}\b")),
]
GENERIC_SECRET_RE = re.compile(
    r"(?i)\b(api[_-]?key|client[_-]?secret|access[_-]?token|refresh[_-]?token|password|passwd|secret)\b"
    r"\s*[:=]\s*[\"']([^\"']{8,})[\"']"
)
DEMO_VALUES = {"changeme", "change_me", "placeholder", "example", "example123", "test12345", "your_token_here"}


@dataclass
class Report:
    blockers: list[str] = field(default_factory=list)
    warnings: list[str] = field(default_factory=list)
    info: list[str] = field(default_factory=list)

    @property
    def status(self) -> str:
        return "BLOCKED" if self.blockers else "READY_FOR_HUMAN_REVIEW"

    def merge(self, other: "Report") -> None:
        self.blockers.extend(other.blockers)
        self.warnings.extend(other.warnings)
        self.info.extend(other.info)


@dataclass
class GitInfo:
    available: bool
    commit: str
    branch: str
    dirty: bool
    status_lines: list[str]


@dataclass
class SourceEntry:
    path: str
    language: str
    sha256: str
    raw_lines: int
    nonempty_lines: int
    included: bool
    reason: str


@dataclass
class SourceBundle:
    entries: list[SourceEntry]
    display_rows: list[str]
    selected_rows: list[str]
    source_files: list[Path]
    raw_lines: int
    nonempty_lines: int
    display_line_count: int
    selected_line_count: int
    expected_pages: int
    truncated_to_front_back: bool
    long_line_count: int


class SoftCopyrightError(RuntimeError):
    pass


def eprint(*args: Any) -> None:
    print(*args, file=sys.stderr)


def sha256_file(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as f:
        for chunk in iter(lambda: f.read(1024 * 1024), b""):
            h.update(chunk)
    return h.hexdigest()


def sha256_bytes(data: bytes) -> str:
    return hashlib.sha256(data).hexdigest()


def load_toml(path: Path) -> dict[str, Any]:
    if not path.exists():
        raise SoftCopyrightError(f"配置文件不存在：{path}")
    with path.open("rb") as f:
        return tomllib.load(f)


def load_json(path: Path) -> dict[str, Any]:
    if not path.exists():
        raise SoftCopyrightError(f"JSON 文件不存在：{path}")
    try:
        return json.loads(path.read_text(encoding="utf-8"))
    except json.JSONDecodeError as exc:
        raise SoftCopyrightError(f"JSON 格式错误：{path}: {exc}") from exc


def ensure_out(out: Path) -> None:
    out.mkdir(parents=True, exist_ok=True)
    (out / ".work").mkdir(parents=True, exist_ok=True)
    (out / ".work" / "logs").mkdir(parents=True, exist_ok=True)


def cfg(cfg_obj: dict[str, Any], section: str, key: str, default: Any = None) -> Any:
    return cfg_obj.get(section, {}).get(key, default)


def posix_rel(path: Path, repo: Path) -> str:
    return path.resolve().relative_to(repo.resolve()).as_posix()


def run_process(
    command: Sequence[str] | str,
    cwd: Path,
    *,
    shell: bool = False,
    timeout: int | None = None,
    env: dict[str, str] | None = None,
) -> subprocess.CompletedProcess[str]:
    return subprocess.run(
        command,
        cwd=str(cwd),
        shell=shell,
        check=False,
        stdout=subprocess.PIPE,
        stderr=subprocess.STDOUT,
        text=True,
        timeout=timeout,
        env=env,
    )


def git_info(repo: Path, ignore_paths: Sequence[Path] | None = None) -> GitInfo:
    check = run_process(["git", "rev-parse", "--is-inside-work-tree"], repo)
    if check.returncode != 0 or check.stdout.strip() != "true":
        return GitInfo(False, "NO_GIT", "NO_GIT", True, ["目录不是 Git 仓库"])
    commit = run_process(["git", "rev-parse", "HEAD"], repo).stdout.strip() or "UNKNOWN"
    branch = run_process(["git", "branch", "--show-current"], repo).stdout.strip() or "DETACHED"
    status = run_process(["git", "status", "--porcelain", "--untracked-files=all"], repo).stdout.splitlines()
    ignore_prefixes: list[str] = []
    for ignored in ignore_paths or []:
        try:
            rel = ignored.resolve().relative_to(repo.resolve()).as_posix().rstrip("/")
        except ValueError:
            continue
        if rel and rel != ".":
            ignore_prefixes.append(rel)
    if ignore_prefixes:
        filtered: list[str] = []
        for line in status:
            raw_path = line[3:].strip() if len(line) > 3 else ""
            candidate_paths = [raw_path]
            if " -> " in raw_path:
                candidate_paths = [part.strip() for part in raw_path.split(" -> ")]
            if any(any(cp == prefix or cp.startswith(prefix + "/") for prefix in ignore_prefixes) for cp in candidate_paths):
                continue
            filtered.append(line)
        status = filtered
    return GitInfo(True, commit, branch, bool(status), status)


def tracked_files(repo: Path, git: GitInfo) -> list[Path]:
    if git.available:
        result = subprocess.run(
            ["git", "ls-files", "-z"], cwd=str(repo), stdout=subprocess.PIPE, stderr=subprocess.PIPE, check=False
        )
        if result.returncode == 0:
            return [repo / p.decode("utf-8", errors="surrogateescape") for p in result.stdout.split(b"\0") if p]
    paths: list[Path] = []
    for root, dirs, files in os.walk(repo):
        dirs[:] = [d for d in dirs if d not in {".git", "node_modules", "vendor", "build", "dist", "target"}]
        for filename in files:
            paths.append(Path(root) / filename)
    return paths


def matches_glob(rel: str, pattern: str) -> bool:
    rel = rel.replace("\\", "/")
    pattern = pattern.replace("\\", "/")
    try:
        if PurePosixPath(rel).match(pattern):
            return True
    except Exception:
        pass
    if fnmatch.fnmatch(rel, pattern):
        return True
    if pattern.startswith("**/") and fnmatch.fnmatch(rel, pattern[3:]):
        return True
    return False


def inside_root(rel: str, root: str) -> bool:
    clean = root.strip().strip("/").replace("\\", "/")
    if not clean or clean == ".":
        return True
    return rel == clean or rel.startswith(clean + "/")


def read_text(path: Path) -> str | None:
    try:
        data = path.read_bytes()
    except OSError:
        return None
    if b"\0" in data[:4096]:
        return None
    for enc in ("utf-8", "utf-8-sig", "gb18030", "big5", "latin-1"):
        try:
            return data.decode(enc)
        except UnicodeDecodeError:
            continue
    return None


def comment_marker(path: Path, rel: str) -> str:
    ext = path.suffix.lower()
    if ext in {".py", ".rb", ".sh", ".bash", ".r"}:
        return f"# FILE: {rel}"
    if ext in {".xml", ".html", ".vue"}:
        return f"<!-- FILE: {rel} -->"
    if ext in {".css", ".scss", ".less"}:
        return f"/* FILE: {rel} */"
    if ext == ".sql":
        return f"-- FILE: {rel}"
    return f"// FILE: {rel}"


def wrap_print_line(line_no: int | None, text: str, width: int) -> tuple[list[str], bool]:
    prefix = f"{line_no:06d} | " if line_no is not None else "       | "
    continuation = "       | "
    usable = max(30, width - len(prefix))
    expanded = text.expandtabs(4).rstrip()
    if not expanded:
        return [prefix], False
    chunks = textwrap.wrap(
        expanded,
        width=usable,
        replace_whitespace=False,
        drop_whitespace=False,
        break_long_words=True,
        break_on_hyphens=False,
    ) or [""]
    rows = [prefix + chunks[0]] + [continuation + chunk for chunk in chunks[1:]]
    return rows, len(chunks) > 1


def collect_source_bundle(repo: Path, config: dict[str, Any], git: GitInfo) -> tuple[SourceBundle, Report]:
    report = Report()
    all_paths = tracked_files(repo, git)
    roots = [str(x).strip() for x in cfg(config, "source", "roots", []) if str(x).strip()]
    extensions = {str(x).lower() for x in cfg(config, "source", "include_extensions", list(DEFAULT_EXTENSIONS))}
    excludes = [str(x) for x in cfg(config, "source", "exclude_globs", [])]
    omit_blank = bool(cfg(config, "source", "omit_blank_lines", True))
    print_width = int(cfg(config, "source", "print_width", 110))
    lines_per_page = int(cfg(config, "source", "lines_per_page", 50))
    if lines_per_page != 50:
        report.warnings.append(f"当前每页代码打印行数为 {lines_per_page}，官方基线通常为每页不少于50行。")
    if cfg(config, "source", "deposit_mode", "general") != "general":
        report.blockers.append("配置要求例外交存或非一般交存，自动化已停止，需人工处理。")

    existing_roots = [root for root in roots if (repo / root).exists()]
    for root in roots:
        if not (repo / root).exists():
            report.warnings.append(f"源代码根目录不存在：{root}")
    if roots and not existing_roots:
        report.blockers.append("配置的 source.roots 均不存在。")

    candidates: list[tuple[Path, str, int, str]] = []
    entries: list[SourceEntry] = []
    for path in all_paths:
        if not path.is_file():
            continue
        try:
            rel = posix_rel(path, repo)
        except ValueError:
            continue
        ext = path.suffix.lower()
        if ext not in extensions:
            continue
        excluded_pattern = next((p for p in excludes if matches_glob(rel, p)), None)
        root_index = -1
        if roots:
            for idx, root in enumerate(roots):
                if inside_root(rel, root):
                    root_index = idx
                    break
        else:
            root_index = 0
        included = excluded_pattern is None and root_index >= 0
        reason = "included"
        if excluded_pattern:
            reason = f"excluded by glob: {excluded_pattern}"
        elif root_index < 0:
            reason = "outside source.roots"
        text = read_text(path)
        raw_count = len(text.splitlines()) if text is not None else 0
        nonempty = sum(1 for line in text.splitlines() if line.strip()) if text is not None else 0
        if text is None:
            included = False
            reason = "binary or unreadable"
        digest = sha256_file(path)
        entry = SourceEntry(
            path=rel,
            language=LANGUAGE_BY_EXT.get(ext, ext.lstrip(".").upper()),
            sha256=digest,
            raw_lines=raw_count,
            nonempty_lines=nonempty,
            included=included,
            reason=reason,
        )
        entries.append(entry)
        if included:
            candidates.append((path, rel, root_index, text or ""))

    candidates.sort(key=lambda item: (item[2], item[1]))
    entries.sort(key=lambda item: item.path)
    if not candidates:
        report.blockers.append("没有找到可纳入的第一方源代码文件。")

    display_rows: list[str] = []
    raw_total = 0
    nonempty_total = 0
    long_line_count = 0
    source_files: list[Path] = []
    for path, rel, _root_index, text in candidates:
        source_files.append(path)
        lines = text.splitlines()
        raw_total += len(lines)
        nonempty_total += sum(1 for line in lines if line.strip())
        marker_rows, marker_long = wrap_print_line(None, comment_marker(path, rel), print_width)
        display_rows.extend(marker_rows)
        long_line_count += int(marker_long)
        for idx, line in enumerate(lines, start=1):
            if omit_blank and not line.strip():
                continue
            rows, was_long = wrap_print_line(idx, line, print_width)
            display_rows.extend(rows)
            long_line_count += int(was_long)

    max_rows = 60 * lines_per_page
    half_rows = 30 * lines_per_page
    truncated = len(display_rows) >= max_rows
    if truncated:
        selected = display_rows[:half_rows] + display_rows[-half_rows:]
    else:
        selected = display_rows[:]
    pages = max(1, (len(selected) + lines_per_page - 1) // lines_per_page) if selected else 0
    if nonempty_total < 500:
        report.warnings.append(f"第一方源程序非空行数仅 {nonempty_total}，请确认软件版本是否已经完成且源码范围是否完整。")
    if long_line_count:
        report.warnings.append(f"检测到 {long_line_count} 个超长代码行，提交版将仅进行视觉换行，不删除内容。")

    bundle = SourceBundle(
        entries=entries,
        display_rows=display_rows,
        selected_rows=selected,
        source_files=source_files,
        raw_lines=raw_total,
        nonempty_lines=nonempty_total,
        display_line_count=len(display_rows),
        selected_line_count=len(selected),
        expected_pages=pages,
        truncated_to_front_back=truncated,
        long_line_count=long_line_count,
    )
    return bundle, report


def write_source_manifest(bundle: SourceBundle, out: Path) -> Path:
    path = out / "02_source-file-manifest.csv"
    with path.open("w", encoding="utf-8-sig", newline="") as f:
        writer = csv.writer(f)
        writer.writerow(["relative_path", "language", "raw_lines", "nonempty_lines", "sha256", "included", "reason"])
        for item in bundle.entries:
            writer.writerow([
                item.path, item.language, item.raw_lines, item.nonempty_lines,
                item.sha256, "yes" if item.included else "no", item.reason,
            ])
    return path


def set_run_font(run: Any, latin: str, east_asia: str, size: float, bold: bool = False) -> None:
    run.font.name = latin
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)


def remove_table_borders(table: Any) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "nil")


def add_page_field(paragraph: Any) -> None:
    paragraph.add_run("第 ")
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    value = OxmlElement("w:t")
    value.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, value, end])
    paragraph.add_run(" 页")


def configure_doc(doc: Any, full_name: str, version: str, *, code: bool = False) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.45 if code else 1.8)
    section.bottom_margin = Cm(1.25 if code else 1.8)
    section.left_margin = Cm(1.35 if code else 2.0)
    section.right_margin = Cm(1.35 if code else 2.0)
    section.header_distance = Cm(0.55)
    section.footer_distance = Cm(0.55)

    header = section.header
    p0 = header.paragraphs[0]
    p0.text = ""
    table = header.add_table(rows=1, cols=2, width=section.page_width - section.left_margin - section.right_margin)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    remove_table_borders(table)
    left, right = table.rows[0].cells
    left.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    right.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    lp = left.paragraphs[0]
    rp = right.paragraphs[0]
    lp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = lp.add_run(f"{full_name} {version}")
    set_run_font(run, "Arial", "宋体", 8.5)
    add_page_field(rp)
    for run in rp.runs:
        set_run_font(run, "Arial", "宋体", 8.5)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(10.5)


def build_source_docx(bundle: SourceBundle, config: dict[str, Any], out: Path) -> tuple[Path, Report]:
    report = Report()
    if Document is None:
        report.blockers.append("缺少 python-docx，无法生成 DOCX。")
        return out / "02_source-code-submission.docx", report
    full_name = str(cfg(config, "software", "full_name", "")).strip()
    version = str(cfg(config, "software", "version", "")).strip()
    if not full_name or not version:
        report.blockers.append("软件全称或版本号缺失，无法生成源程序页眉。")
        return out / "02_source-code-submission.docx", report
    if not bundle.selected_rows:
        report.blockers.append("没有源程序打印行，无法生成源程序材料。")
        return out / "02_source-code-submission.docx", report

    lines_per_page = int(cfg(config, "source", "lines_per_page", 50))
    doc = Document()
    configure_doc(doc, full_name, version, code=True)
    style = doc.styles.add_style("SourceCodeLine", 1)  # WD_STYLE_TYPE.PARAGRAPH = 1
    style.font.name = "Courier New"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "等线")
    style.font.size = Pt(7.2)
    pf = style.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = Pt(8.25)

    for idx, row in enumerate(bundle.selected_rows):
        p = doc.add_paragraph(style="SourceCodeLine")
        p.paragraph_format.keep_together = False
        p.paragraph_format.keep_with_next = False
        run = p.add_run(row)
        set_run_font(run, "Courier New", "等线", 7.2)
        if (idx + 1) % lines_per_page == 0 and idx + 1 < len(bundle.selected_rows):
            p.add_run().add_break(WD_BREAK.PAGE)

    path = out / "02_source-code-submission.docx"
    doc.save(path)
    return path, report


def find_libreoffice() -> str | None:
    for name in ("libreoffice", "soffice"):
        found = shutil.which(name)
        if found:
            return found
    return None


def convert_docx_to_pdf(docx_path: Path, target_pdf: Path) -> tuple[bool, str]:
    lo = find_libreoffice()
    if not lo:
        return False, "未找到 LibreOffice"
    target_pdf.parent.mkdir(parents=True, exist_ok=True)
    with tempfile.TemporaryDirectory(prefix="softcopyright-lo-") as tmp:
        tmp_path = Path(tmp)
        env = os.environ.copy()
        env["HOME"] = str(tmp_path / "home")
        Path(env["HOME"]).mkdir(parents=True, exist_ok=True)
        profile = tmp_path / "lo-profile"
        profile.mkdir(parents=True, exist_ok=True)
        profile_uri = profile.resolve().as_uri()
        result = run_process(
            [
                lo,
                f"-env:UserInstallation={profile_uri}",
                "--headless",
                "--norestore",
                "--nodefault",
                "--nolockcheck",
                "--convert-to",
                "pdf",
                "--outdir",
                str(tmp_path),
                str(docx_path),
            ],
            docx_path.parent,
            timeout=180,
            env=env,
        )
        generated = tmp_path / (docx_path.stem + ".pdf")
        if result.returncode != 0 or not generated.exists():
            return False, result.stdout.strip() or "LibreOffice 转换失败"
        shutil.copy2(generated, target_pdf)
    return True, "ok"


def pdf_page_count(path: Path) -> int | None:
    if PdfReader is None or not path.exists():
        return None
    try:
        return len(PdfReader(str(path)).pages)
    except Exception:
        return None


def add_image(doc: Any, image_path: Path, caption: str, report: Report) -> None:
    if not image_path.exists():
        report.blockers.append(f"说明书图片不存在：{image_path}")
        p = doc.add_paragraph(f"[缺失图片：{caption}]")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return
    try:
        max_w = 6.2
        max_h = 6.8
        width = Inches(max_w)
        if Image is not None:
            with Image.open(image_path) as im:
                w, h = im.size
            if w and h:
                ratio = w / h
                width_in = min(max_w, max_h * ratio)
                width = Inches(max(1.0, width_in))
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(image_path), width=width)
        cp = doc.add_paragraph(f"图：{caption}")
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in cp.runs:
            set_run_font(run, "Arial", "宋体", 9)
    except Exception as exc:
        report.blockers.append(f"插入图片失败：{image_path}: {exc}")


def is_table_separator(line: str) -> bool:
    stripped = line.strip().strip("|")
    cells = [c.strip() for c in stripped.split("|")]
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", c or "") for c in cells)


def split_table_row(line: str) -> list[str]:
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def add_markdown_table(doc: Any, rows: list[list[str]]) -> None:
    cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for r_idx, row in enumerate(rows):
        for c_idx in range(cols):
            text = row[c_idx] if c_idx < len(row) else ""
            cell = table.cell(r_idx, c_idx)
            p = cell.paragraphs[0]
            run = p.add_run(text)
            set_run_font(run, "Arial", "宋体", 9.5, bold=(r_idx == 0))


def markdown_to_docx(manual_path: Path, config: dict[str, Any], out_path: Path) -> Report:
    report = Report()
    if Document is None:
        report.blockers.append("缺少 python-docx，无法生成说明书 DOCX。")
        return report
    if not manual_path.exists():
        report.blockers.append(f"说明书 Markdown 不存在：{manual_path}")
        return report
    full_name = str(cfg(config, "software", "full_name", "")).strip()
    version = str(cfg(config, "software", "version", "")).strip()
    raw = manual_path.read_text(encoding="utf-8")
    raw = raw.replace("{{SOFTWARE_FULL_NAME}}", full_name).replace("{{VERSION}}", version)
    placeholders = sorted(set(m.group(0) for m in PLACEHOLDER_RE.finditer(raw)))
    if placeholders:
        report.blockers.append("说明书仍含占位符：" + "、".join(placeholders[:10]))

    doc = Document()
    configure_doc(doc, full_name, version, code=False)
    for style_name, size in (("Title", 18), ("Heading 1", 15), ("Heading 2", 13), ("Heading 3", 11.5)):
        style = doc.styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor(0, 0, 0)
    normal = doc.styles["Normal"]
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(4)

    lines = raw.splitlines()
    i = 0
    in_code = False
    code_lines: list[str] = []
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        if stripped.startswith("```"):
            if not in_code:
                in_code = True
                code_lines = []
            else:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(0.5)
                for idx, code_line in enumerate(code_lines):
                    run = p.add_run(code_line + ("\n" if idx + 1 < len(code_lines) else ""))
                    set_run_font(run, "Courier New", "等线", 8.5)
                in_code = False
            i += 1
            continue
        if in_code:
            code_lines.append(line)
            i += 1
            continue
        if not stripped:
            i += 1
            continue
        if stripped == "[[PAGE_BREAK]]":
            doc.add_page_break()
            i += 1
            continue
        image_match = re.fullmatch(r"!\[([^]]*)\]\(([^)]+)\)", stripped)
        if image_match:
            caption, ref = image_match.groups()
            image_path = (manual_path.parent / ref).resolve()
            add_image(doc, image_path, caption or image_path.stem, report)
            i += 1
            continue
        if stripped.startswith("|") and i + 1 < len(lines) and is_table_separator(lines[i + 1]):
            rows = [split_table_row(line)]
            i += 2
            while i < len(lines) and lines[i].strip().startswith("|"):
                rows.append(split_table_row(lines[i]))
                i += 1
            add_markdown_table(doc, rows)
            continue
        heading = re.match(r"^(#{1,4})\s+(.+)$", stripped)
        if heading:
            level = len(heading.group(1))
            text = heading.group(2).strip()
            if level == 1:
                p = doc.add_paragraph(style="Title")
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                p = doc.add_paragraph(style=f"Heading {min(level - 1, 3)}")
            run = p.add_run(text)
            set_run_font(run, "Arial", "黑体", 18 if level == 1 else (15 if level == 2 else 13), bold=True)
            i += 1
            continue
        bullet = re.match(r"^[-*]\s+(.+)$", stripped)
        if bullet:
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(bullet.group(1))
            set_run_font(run, "Arial", "宋体", 10.5)
            i += 1
            continue
        numbered = re.match(r"^\d+[.)]\s+(.+)$", stripped)
        if numbered:
            p = doc.add_paragraph(style="List Number")
            run = p.add_run(numbered.group(1))
            set_run_font(run, "Arial", "宋体", 10.5)
            i += 1
            continue

        paragraph_lines = [stripped]
        i += 1
        while i < len(lines):
            nxt = lines[i].strip()
            if not nxt:
                break
            if (
                nxt.startswith("#") or nxt.startswith("- ") or nxt.startswith("* ")
                or re.match(r"^\d+[.)]\s+", nxt) or nxt.startswith("![")
                or nxt == "[[PAGE_BREAK]]" or nxt.startswith("```") or nxt.startswith("|")
            ):
                break
            paragraph_lines.append(nxt)
            i += 1
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0.74)
        run = p.add_run(" ".join(paragraph_lines))
        set_run_font(run, "Arial", "宋体", 10.5)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    return report


def split_manual_pdf(full_pdf: Path, submission_pdf: Path) -> tuple[int | None, str]:
    if PdfReader is None or PdfWriter is None:
        return None, "缺少 pypdf"
    try:
        reader = PdfReader(str(full_pdf))
        count = len(reader.pages)
        writer = PdfWriter()
        if count > 60:
            for page in reader.pages[:30]:
                writer.add_page(page)
            for page in reader.pages[-30:]:
                writer.add_page(page)
        else:
            for page in reader.pages:
                writer.add_page(page)
        with submission_pdf.open("wb") as f:
            writer.write(f)
        return count, "front30+back30" if count > 60 else "all"
    except Exception as exc:
        return None, str(exc)


def validate_required_facts(config: dict[str, Any]) -> Report:
    report = Report()
    software = config.get("software", {})
    required = {
        "full_name": "软件全称",
        "short_name": "软件简称",
        "version": "版本号",
        "completion_date": "开发完成日期",
        "publication_status": "发表状态",
        "manual_type": "说明书类型",
    }
    for key, label in required.items():
        value = str(software.get(key, "")).strip()
        if not value:
            report.blockers.append(f"缺少必填字段：{label}（software.{key}）")
    status = str(software.get("publication_status", "")).strip()
    if status not in {"unpublished", "published"}:
        report.blockers.append("software.publication_status 必须为 unpublished 或 published。")
    if status == "published" and not str(software.get("first_publication_date", "")).strip():
        report.blockers.append("软件标记为已发表，但未填写首次发表日期。")
    manual_type = str(software.get("manual_type", "")).strip()
    if manual_type not in {"user-manual", "design-description"}:
        report.blockers.append("software.manual_type 必须为 user-manual 或 design-description。")
    completion = str(software.get("completion_date", "")).strip()
    if completion:
        try:
            date_value = dt.date.fromisoformat(completion)
            if date_value > dt.date.today():
                report.blockers.append(f"开发完成日期 {completion} 在未来。")
        except ValueError:
            report.blockers.append("开发完成日期必须使用 YYYY-MM-DD 格式。")
    publication = str(software.get("first_publication_date", "")).strip()
    if publication:
        try:
            pub_date = dt.date.fromisoformat(publication)
            if pub_date > dt.date.today():
                report.blockers.append(f"首次发表日期 {publication} 在未来。")
        except ValueError:
            report.blockers.append("首次发表日期必须使用 YYYY-MM-DD 格式。")
    if cfg(config, "source", "deposit_mode", "general") != "general":
        report.blockers.append("非一般交存必须由人工依据官方要求处理。")
    age = (dt.date.today() - RULES_CHECKED).days
    if age > 180:
        report.warnings.append(f"内置规则已超过180天未核验（最后核验 {RULES_CHECKED.isoformat()}），请先核对官方最新要求。")
    return report


def validate_git(config: dict[str, Any], git: GitInfo) -> Report:
    report = Report()
    require_clean = bool(cfg(config, "build", "require_clean_git", True))
    if not git.available:
        report.blockers.append("项目不是 Git 仓库，无法将材料绑定到可复现的代码快照。")
    elif git.dirty and require_clean:
        preview = "；".join(git.status_lines[:8])
        report.blockers.append(f"Git 工作区存在未提交修改：{preview}")
    elif git.dirty:
        report.warnings.append("Git 工作区存在未提交修改，材料的可复现性降低。")
    return report


def execute_configured_commands(repo: Path, config: dict[str, Any], out: Path) -> Report:
    report = Report()
    for key, label in (("build_command", "构建"), ("test_command", "测试"), ("screenshot_command", "截图")):
        command = str(cfg(config, "build", key, "")).strip()
        if not command:
            if key in {"build_command", "test_command"}:
                report.warnings.append(f"未配置{label}命令，无法自动验证当前版本。")
            continue
        log_path = out / ".work" / "logs" / f"{key}.log"
        try:
            result = run_process(command, repo, shell=True, timeout=1800)
            log_path.write_text(result.stdout, encoding="utf-8", errors="replace")
            if result.returncode != 0:
                report.blockers.append(f"{label}命令失败，退出码 {result.returncode}，详见 {log_path}")
            else:
                report.info.append(f"{label}命令执行成功。")
        except subprocess.TimeoutExpired:
            report.blockers.append(f"{label}命令执行超时。")
    return report


def scan_secrets(repo: Path, source_files: Iterable[Path], all_paths: Iterable[Path]) -> tuple[list[dict[str, Any]], Report]:
    report = Report()
    findings: list[dict[str, Any]] = []
    for path in all_paths:
        if not path.is_file():
            continue
        try:
            rel = posix_rel(path, repo)
        except ValueError:
            continue
        lower_name = path.name.lower()
        if path.suffix.lower() in SENSITIVE_SUFFIXES or lower_name == ".env" or lower_name.startswith(".env."):
            findings.append({"severity": "high", "type": "sensitive-file", "path": rel, "line": None})
            report.blockers.append(f"检测到敏感文件：{rel}")

    for path in source_files:
        text = read_text(path)
        if text is None:
            continue
        rel = posix_rel(path, repo)
        for line_no, line in enumerate(text.splitlines(), start=1):
            for kind, pattern in HIGH_SECRET_PATTERNS:
                if pattern.search(line):
                    findings.append({"severity": "high", "type": kind, "path": rel, "line": line_no})
                    report.blockers.append(f"检测到疑似高风险凭据：{rel}:{line_no} ({kind})")
            generic = GENERIC_SECRET_RE.search(line)
            if generic:
                value = generic.group(2).strip().lower()
                if value not in DEMO_VALUES and not any(x in value for x in ("example", "placeholder", "dummy", "test")):
                    findings.append({"severity": "medium", "type": "generic-secret", "path": rel, "line": line_no})
                    report.warnings.append(f"检测到疑似硬编码秘密，请人工检查：{rel}:{line_no}")
    return findings, report


def validate_evidence(repo: Path, evidence: dict[str, Any], manual_type: str) -> Report:
    report = Report()
    features = evidence.get("features")
    if not isinstance(features, list) or not features:
        report.blockers.append("feature-evidence.json 未包含有效的 features 列表。")
        return report
    names: set[str] = set()
    for idx, feature in enumerate(features, start=1):
        if not isinstance(feature, dict):
            report.blockers.append(f"第 {idx} 个功能证据不是对象。")
            continue
        name = str(feature.get("name", "")).strip()
        status = str(feature.get("status", "")).strip()
        if not name:
            report.blockers.append(f"第 {idx} 个功能缺少 name。")
            continue
        if name in names:
            report.blockers.append(f"功能名称重复：{name}")
        names.add(name)
        if status not in {"verified", "needs_review", "excluded"}:
            report.blockers.append(f"功能 {name} 的 status 无效。")
            continue
        if status == "needs_review":
            report.warnings.append(f"功能证据待复核：{name}")
        if status == "excluded":
            continue
        code_paths = feature.get("code_paths", [])
        screenshots = feature.get("screenshots", [])
        steps = feature.get("operation_steps", [])
        if not isinstance(code_paths, list) or not code_paths:
            report.blockers.append(f"功能 {name} 缺少代码证据路径。")
        for ref in code_paths if isinstance(code_paths, list) else []:
            p = (repo / str(ref)).resolve()
            if not p.exists():
                report.blockers.append(f"功能 {name} 引用的代码路径不存在：{ref}")
        if status == "verified" and manual_type == "user-manual" and (not isinstance(screenshots, list) or not screenshots):
            report.warnings.append(f"已核验的界面功能 {name} 没有关联截图，请确认是否属于无界面功能。")
        for ref in screenshots if isinstance(screenshots, list) else []:
            p = (repo / str(ref)).resolve()
            if not p.exists():
                report.blockers.append(f"功能 {name} 引用的截图不存在：{ref}")
            elif p.suffix.lower() not in {".png", ".jpg", ".jpeg", ".webp"}:
                report.warnings.append(f"功能 {name} 的截图格式不常见：{ref}")
        if status == "verified" and (not isinstance(steps, list) or not steps):
            report.warnings.append(f"功能 {name} 缺少可复现操作步骤。")
    return report


def detect_project_metadata(repo: Path, all_paths: list[Path]) -> dict[str, Any]:
    detected: dict[str, Any] = {"markers": [], "names": [], "versions": [], "package_ids": [], "platforms": []}
    by_rel = {posix_rel(p, repo): p for p in all_paths if p.is_file() and p.exists()}

    def add(bucket: str, value: Any) -> None:
        if value is None:
            return
        s = str(value).strip()
        if s and s not in detected[bucket]:
            detected[bucket].append(s)

    for rel, path in by_rel.items():
        name = path.name
        if name == "package.json":
            try:
                data = json.loads(path.read_text(encoding="utf-8"))
                add("names", data.get("productName") or data.get("name"))
                add("versions", data.get("version"))
                detected["markers"].append(rel)
            except Exception:
                pass
        elif name == "pubspec.yaml":
            text = read_text(path) or ""
            m_name = re.search(r"(?m)^name:\s*([^#\n]+)", text)
            m_version = re.search(r"(?m)^version:\s*([^#\n]+)", text)
            add("names", m_name.group(1) if m_name else None)
            add("versions", m_version.group(1).split("+")[0] if m_version else None)
            add("platforms", "Flutter")
            detected["markers"].append(rel)
        elif name in {"build.gradle", "build.gradle.kts"}:
            text = read_text(path) or ""
            app_id = re.search(r"applicationId\s*[=( ]\s*[\"']([^\"']+)", text)
            version = re.search(r"versionName\s*[=( ]\s*[\"']([^\"']+)", text)
            add("package_ids", app_id.group(1) if app_id else None)
            add("versions", version.group(1) if version else None)
            if "com.android.application" in text or "applicationId" in text:
                add("platforms", "Android")
            detected["markers"].append(rel)
        elif name == "AndroidManifest.xml":
            text = read_text(path) or ""
            pkg = re.search(r"\bpackage=[\"']([^\"']+)", text)
            add("package_ids", pkg.group(1) if pkg else None)
            add("platforms", "Android")
            detected["markers"].append(rel)
        elif name == "Info.plist":
            try:
                with path.open("rb") as f:
                    data = plistlib.load(f)
                add("names", data.get("CFBundleDisplayName") or data.get("CFBundleName"))
                add("versions", data.get("CFBundleShortVersionString"))
                add("package_ids", data.get("CFBundleIdentifier"))
                add("platforms", "iOS/macOS")
                detected["markers"].append(rel)
            except Exception:
                pass
        elif name == "pyproject.toml":
            try:
                with path.open("rb") as f:
                    data = tomllib.load(f)
                project = data.get("project", {})
                add("names", project.get("name"))
                add("versions", project.get("version"))
                add("platforms", "Python")
                detected["markers"].append(rel)
            except Exception:
                pass
        elif name == "Cargo.toml":
            try:
                with path.open("rb") as f:
                    data = tomllib.load(f)
                package = data.get("package", {})
                add("names", package.get("name"))
                add("versions", package.get("version"))
                add("platforms", "Rust")
                detected["markers"].append(rel)
            except Exception:
                pass
        elif name == "go.mod":
            text = read_text(path) or ""
            module = re.search(r"(?m)^module\s+(.+)$", text)
            add("package_ids", module.group(1).strip() if module else None)
            add("platforms", "Go")
            detected["markers"].append(rel)
    return detected


def detect_dependencies(repo: Path, all_paths: list[Path]) -> tuple[list[dict[str, str]], list[str]]:
    dependencies: list[dict[str, str]] = []
    manifests: list[str] = []
    for path in all_paths:
        if not path.is_file() or not path.exists():
            continue
        rel = posix_rel(path, repo)
        name = path.name
        try:
            if name == "package.json":
                data = json.loads(path.read_text(encoding="utf-8"))
                manifests.append(rel)
                for group in ("dependencies", "devDependencies", "peerDependencies"):
                    for dep, ver in (data.get(group, {}) or {}).items():
                        dependencies.append({"ecosystem": "npm", "name": dep, "version": str(ver), "source": rel})
            elif name == "requirements.txt":
                manifests.append(rel)
                for line in (read_text(path) or "").splitlines():
                    line = line.strip()
                    if not line or line.startswith("#") or line.startswith("-"):
                        continue
                    m = re.match(r"([A-Za-z0-9_.-]+)\s*([<>=!~].+)?", line)
                    if m:
                        dependencies.append({"ecosystem": "pip", "name": m.group(1), "version": (m.group(2) or "").strip(), "source": rel})
            elif name == "go.mod":
                manifests.append(rel)
                for line in (read_text(path) or "").splitlines():
                    m = re.match(r"\s*([A-Za-z0-9_.\-/]+)\s+(v[^\s]+)", line)
                    if m and not line.strip().startswith("module"):
                        dependencies.append({"ecosystem": "go", "name": m.group(1), "version": m.group(2), "source": rel})
            elif name == "Cargo.toml":
                manifests.append(rel)
                with path.open("rb") as f:
                    data = tomllib.load(f)
                for group in ("dependencies", "dev-dependencies", "build-dependencies"):
                    for dep, spec in (data.get(group, {}) or {}).items():
                        dependencies.append({"ecosystem": "cargo", "name": dep, "version": str(spec), "source": rel})
            elif name == "pyproject.toml":
                manifests.append(rel)
                with path.open("rb") as f:
                    data = tomllib.load(f)
                for spec in data.get("project", {}).get("dependencies", []) or []:
                    m = re.match(r"([A-Za-z0-9_.-]+)(.*)", str(spec))
                    if m:
                        dependencies.append({"ecosystem": "pip", "name": m.group(1), "version": m.group(2).strip(), "source": rel})
            elif name in {"build.gradle", "build.gradle.kts"}:
                manifests.append(rel)
                text = read_text(path) or ""
                for m in re.finditer(r"(?:implementation|api|compileOnly|runtimeOnly|testImplementation)\s*\(?[\"']([^\"']+)[\"']", text):
                    coordinate = m.group(1)
                    parts = coordinate.split(":")
                    dep_name = ":".join(parts[:2]) if len(parts) >= 2 else coordinate
                    dep_ver = parts[2] if len(parts) >= 3 else ""
                    dependencies.append({"ecosystem": "gradle", "name": dep_name, "version": dep_ver, "source": rel})
            elif name == "pubspec.yaml":
                manifests.append(rel)
                text = read_text(path) or ""
                in_deps = False
                for line in text.splitlines():
                    if re.match(r"^(dependencies|dev_dependencies):\s*$", line):
                        in_deps = True
                        continue
                    if in_deps and line and not line.startswith((" ", "\t")):
                        in_deps = False
                    if in_deps:
                        m = re.match(r"\s{2,}([A-Za-z0-9_.-]+):\s*(.*)$", line)
                        if m and m.group(1) != "flutter":
                            dependencies.append({"ecosystem": "pub", "name": m.group(1), "version": m.group(2).strip(), "source": rel})
        except Exception:
            continue
    unique: dict[tuple[str, str, str], dict[str, str]] = {}
    for dep in dependencies:
        key = (dep["ecosystem"], dep["name"], dep["version"])
        unique[key] = dep
    return sorted(unique.values(), key=lambda x: (x["ecosystem"], x["name"])), sorted(set(manifests))


def build_metadata_draft(
    config: dict[str, Any], git: GitInfo, bundle: SourceBundle, evidence: dict[str, Any], detected: dict[str, Any]
) -> dict[str, Any]:
    software = config.get("software", {})
    languages = sorted({entry.language for entry in bundle.entries if entry.included})
    verified_features = [
        f for f in evidence.get("features", [])
        if isinstance(f, dict) and f.get("status") == "verified"
    ]
    functions = [str(f.get("summary") or f.get("name", "")).strip() for f in verified_features]
    functions = [x for x in functions if x]
    draft = {
        "tool_version": TOOL_VERSION,
        "generated_at": dt.datetime.now(dt.timezone.utc).astimezone().isoformat(timespec="seconds"),
        "software": {
            "full_name": software.get("full_name", ""),
            "short_name": software.get("short_name", ""),
            "version": software.get("version", ""),
            "completion_date": software.get("completion_date", ""),
            "publication_status": software.get("publication_status", ""),
            "first_publication_date": software.get("first_publication_date", ""),
            "copyright_owner": software.get("copyright_owner", ""),
            "manual_type": software.get("manual_type", ""),
            "platforms": software.get("platforms", []),
            "app_display_name": software.get("app_display_name", ""),
            "package_id": software.get("package_id", ""),
        },
        "technical_draft": {
            "hardware_environment": software.get("hardware_environment", ""),
            "software_environment": software.get("software_environment", ""),
            "programming_languages": languages,
            "source_raw_lines": bundle.raw_lines,
            "source_nonempty_lines": bundle.nonempty_lines,
            "source_print_lines": bundle.display_line_count,
            "source_submission_pages": bundle.expected_pages,
            "main_functions": functions,
            "technical_characteristics": software.get("technical_characteristics", ""),
        },
        "repository": {
            "git_commit": git.commit,
            "git_branch": git.branch,
            "git_dirty": git.dirty,
        },
        "detected_project_metadata": detected,
        "needs_human_review": [
            "软件全称和简称",
            "版本号",
            "开发完成日期",
            "发表状态和首次发表日期",
            "著作权人",
            "开发方式和权利取得方式",
            "源代码范围是否均为申请人有权处分的第一方代码",
            "运行环境、主要功能和技术特点",
        ],
    }
    return draft


def write_metadata_files(metadata: dict[str, Any], out: Path) -> tuple[Path, Path]:
    json_path = out / "01_software-info-draft.json"
    md_path = out / "01_software-info-draft.md"
    json_path.write_text(json.dumps(metadata, ensure_ascii=False, indent=2), encoding="utf-8")
    sw = metadata["software"]
    tech = metadata["technical_draft"]
    repo = metadata["repository"]
    lines = [
        "# 软件基本信息与申请表技术字段草稿",
        "",
        f"- 软件全称：{sw.get('full_name', '')}",
        f"- 软件简称：{sw.get('short_name', '')}",
        f"- 版本号：{sw.get('version', '')}",
        f"- 开发完成日期：{sw.get('completion_date', '')}（须人工确认）",
        f"- 发表状态：{sw.get('publication_status', '')}（须人工确认）",
        f"- 首次发表日期：{sw.get('first_publication_date', '')}（如适用，须人工确认）",
        f"- 著作权人：{sw.get('copyright_owner', '')}（仅复制人工提供值）",
        f"- 平台：{'、'.join(sw.get('platforms', []))}",
        f"- 应用显示名称：{sw.get('app_display_name', '')}",
        f"- 包名/标识：{sw.get('package_id', '')}",
        "",
        "## 技术字段草稿",
        "",
        f"- 硬件环境：{tech.get('hardware_environment', '')}",
        f"- 软件环境：{tech.get('software_environment', '')}",
        f"- 编程语言：{'、'.join(tech.get('programming_languages', []))}",
        f"- 源程序原始行数：{tech.get('source_raw_lines', 0)}",
        f"- 源程序非空行数：{tech.get('source_nonempty_lines', 0)}",
        f"- 规范化打印行数：{tech.get('source_print_lines', 0)}",
        f"- 源程序提交版页数：{tech.get('source_submission_pages', 0)}",
        f"- 技术特点：{tech.get('technical_characteristics', '')}",
        "",
        "### 主要功能",
        "",
    ]
    functions = tech.get("main_functions", [])
    lines.extend([f"- {x}" for x in functions] or ["- 未发现已核验功能，需补充 feature-evidence.json。"])
    lines.extend([
        "",
        "## 代码快照",
        "",
        f"- Git commit：`{repo.get('git_commit', '')}`",
        f"- Git branch：`{repo.get('git_branch', '')}`",
        f"- 工作区是否有未提交修改：{repo.get('git_dirty', False)}",
        "",
        "> 本文件是技术字段草稿，不替代申请人对法律事实的确认。",
    ])
    md_path.write_text("\n".join(lines) + "\n", encoding="utf-8")
    return json_path, md_path


def write_dependency_report(dependencies: list[dict[str, str]], manifests: list[str], bundle: SourceBundle, out: Path) -> Path:
    path = out / "05_dependency-and-ip-report.md"
    excluded = [entry for entry in bundle.entries if not entry.included]
    lines = [
        "# 第三方依赖与知识产权风险报告",
        "",
        "> 本报告用于材料整理，不构成正式法律意见或完整开源许可证审计。",
        "",
        "## 检测到的依赖清单文件",
        "",
    ]
    lines.extend([f"- `{x}`" for x in manifests] or ["- 未检测到常见依赖清单文件。"])
    lines.extend(["", "## 检测到的第三方依赖", ""])
    if dependencies:
        lines.append("| 生态 | 依赖 | 版本 | 来源文件 |")
        lines.append("|---|---|---|---|")
        for dep in dependencies:
            lines.append(f"| {dep['ecosystem']} | {dep['name']} | {dep['version']} | `{dep['source']}` |")
    else:
        lines.append("- 未从常见清单中解析到依赖；仍需人工检查 SDK、复制代码和子模块。")
    lines.extend(["", "## 被排除的源码候选", ""])
    if excluded:
        for entry in excluded[:200]:
            lines.append(f"- `{entry.path}`：{entry.reason}")
        if len(excluded) > 200:
            lines.append(f"- 其余 {len(excluded) - 200} 项见 `02_source-file-manifest.csv`。")
    else:
        lines.append("- 无。")
    lines.extend([
        "",
        "## 人工复核要点",
        "",
        "- 外包、合作开发、员工职务开发和个人转公司的权属文件是否齐全；",
        "- 第三方 SDK、开源组件和复制代码是否从源程序提交版中排除；",
        "- AI 辅助生成或代码生成器产出的代码是否具有可处分权利并已按项目规则审查；",
        "- 依赖许可证是否允许当前分发和商业使用；",
        "- 是否存在需要例外交存或保密处理的商业秘密。",
    ])
    path.write_text("\n".join(lines) + "\n", encoding="utf-8")
    return path


def copy_normalized_evidence(evidence: dict[str, Any], out: Path) -> Path:
    path = out / "04_feature-evidence.json"
    path.write_text(json.dumps(evidence, ensure_ascii=False, indent=2), encoding="utf-8")
    return path


def write_secret_findings(findings: list[dict[str, Any]], out: Path) -> None:
    (out / ".work" / "secret-findings.json").write_text(
        json.dumps(findings, ensure_ascii=False, indent=2), encoding="utf-8"
    )


def build_inspection(repo: Path, config_path: Path, evidence_path: Path, out: Path, *, run_commands: bool) -> tuple[dict[str, Any], Report]:
    ensure_out(out)
    config = load_toml(config_path)
    evidence = load_json(evidence_path)
    report = Report()
    report.merge(validate_required_facts(config))
    git = git_info(repo, [out])
    report.merge(validate_git(config, git))
    all_paths = tracked_files(repo, git)
    bundle, source_report = collect_source_bundle(repo, config, git)
    report.merge(source_report)
    write_source_manifest(bundle, out)
    report.merge(validate_evidence(repo, evidence, str(cfg(config, "software", "manual_type", ""))))
    if run_commands:
        report.merge(execute_configured_commands(repo, config, out))
    findings, secret_report = scan_secrets(repo, bundle.source_files, all_paths)
    report.merge(secret_report)
    write_secret_findings(findings, out)
    detected = detect_project_metadata(repo, all_paths)
    metadata = build_metadata_draft(config, git, bundle, evidence, detected)
    write_metadata_files(metadata, out)
    dependencies, manifests = detect_dependencies(repo, all_paths)
    write_dependency_report(dependencies, manifests, bundle, out)
    copy_normalized_evidence(evidence, out)

    configured_package = str(cfg(config, "software", "package_id", "")).strip()
    if configured_package and detected.get("package_ids") and configured_package not in detected["package_ids"]:
        report.warnings.append(
            f"配置包名 {configured_package} 与仓库检测值 {detected['package_ids']} 不完全一致，请人工核对。"
        )
    configured_display = str(cfg(config, "software", "app_display_name", "")).strip()
    if configured_display and detected.get("names") and configured_display not in detected["names"]:
        report.warnings.append(
            f"配置应用显示名称 {configured_display} 与仓库检测名称 {detected['names']} 不完全一致，请人工核对。"
        )

    state = {
        "tool_version": TOOL_VERSION,
        "generated_at": dt.datetime.now(dt.timezone.utc).astimezone().isoformat(timespec="seconds"),
        "repo": str(repo.resolve()),
        "config": str(config_path.resolve()),
        "evidence": str(evidence_path.resolve()),
        "git": git.__dict__,
        "source": {
            "raw_lines": bundle.raw_lines,
            "nonempty_lines": bundle.nonempty_lines,
            "display_lines": bundle.display_line_count,
            "selected_lines": bundle.selected_line_count,
            "expected_pages": bundle.expected_pages,
            "front_back_selection": bundle.truncated_to_front_back,
            "long_line_count": bundle.long_line_count,
        },
        "detected": detected,
        "report": report.__dict__,
    }
    (out / ".work" / "inspection.json").write_text(json.dumps(state, ensure_ascii=False, indent=2), encoding="utf-8")
    return state, report


def write_validation_report(report: Report, config: dict[str, Any], git: GitInfo, out: Path) -> Path:
    path = out / "06_consistency-validation.md"
    full_name = str(cfg(config, "software", "full_name", ""))
    version = str(cfg(config, "software", "version", ""))
    lines = [
        "# 材料一致性与风险校验报告",
        "",
        f"- 状态：**{report.status}**",
        f"- 软件：{full_name} {version}",
        f"- Git commit：`{git.commit}`",
        f"- 生成时间：{dt.datetime.now().astimezone().isoformat(timespec='seconds')}",
        "",
        "## 阻断项",
        "",
    ]
    lines.extend([f"- {x}" for x in report.blockers] or ["- 无。"])
    lines.extend(["", "## 警告", ""])
    lines.extend([f"- {x}" for x in report.warnings] or ["- 无。"])
    lines.extend(["", "## 已执行检查", ""])
    lines.extend([f"- {x}" for x in report.info] or ["- 已执行配置、Git、源码、证据、敏感信息和输出文件检查。"])
    path.write_text("\n".join(lines) + "\n", encoding="utf-8")
    return path


def write_human_checklist(config: dict[str, Any], report: Report, out: Path) -> Path:
    path = out / "07_human-review-checklist.md"
    owner = str(cfg(config, "software", "copyright_owner", "")).strip()
    lines = [
        "# 软著技术材料人工复核清单",
        "",
        "在提交申请前，由申请负责人逐项确认并留存确认记录：",
        "",
        "- [ ] 软件全称、简称和版本号与正式申请表完全一致；",
        "- [ ] 开发完成日期是真实日期，不是 Git 时间或自动推断值；",
        "- [ ] 首次发表状态和日期真实、准确；",
        f"- [ ] 著作权人已经确认：{owner or '尚未填写'}；",
        "- [ ] 开发方式、权利取得方式及外包/合作/员工开发权属已经确认；",
        "- [ ] `02_source-file-manifest.csv` 中纳入的文件均为有权处分的第一方代码；",
        "- [ ] 源程序材料不含第三方库、生成代码、密钥、证书、生产配置或用户数据；",
        "- [ ] 说明书中的每项功能在当前版本真实存在并可复现；",
        "- [ ] 所有截图来自真实运行版本，且已去除个人信息和生产数据；",
        "- [ ] DOCX/PDF 已逐页视觉检查，无乱码、截断、重叠、黑块和错误页眉；",
        "- [ ] 如使用 PDF 提交版，页数和前30页/后30页选择已人工核对；",
        "- [ ] 已重新核对当前受理渠道的最新文件格式和上传要求；",
        "- [ ] 公司营业执照、签章、授权委托和权属证明由主体材料流程另行准备。",
        "",
        f"当前自动校验状态：**{report.status}**",
    ]
    if report.blockers:
        lines.extend(["", "## 必须先解决的阻断项", ""] + [f"- [ ] {x}" for x in report.blockers])
    path.write_text("\n".join(lines) + "\n", encoding="utf-8")
    return path


def write_run_summary(config: dict[str, Any], git: GitInfo, report: Report, out: Path) -> Path:
    path = out / "00_run-summary.md"
    full_name = str(cfg(config, "software", "full_name", ""))
    version = str(cfg(config, "software", "version", ""))
    generated = sorted(
        p.name for p in out.iterdir()
        if p.is_file() and p.name not in {"00_run-summary.md", "software-copyright-technical-package.zip"}
    )
    lines = [
        "# 软件著作权技术材料生成结果",
        "",
        f"- 状态：**{report.status}**",
        f"- 版本：{full_name} {version}",
        f"- 代码快照：`{git.commit}`",
        f"- 分支：`{git.branch}`",
        f"- 生成时间：{dt.datetime.now().astimezone().isoformat(timespec='seconds')}",
        "",
        "## 已生成文件",
        "",
    ]
    lines.extend([f"- `{x}`" for x in generated] or ["- 暂无。"])
    lines.extend(["", "## 警告", ""])
    lines.extend([f"- {x}" for x in report.warnings] or ["- 无。"])
    lines.extend(["", "## 必须人工确认", ""])
    lines.extend([
        "- 软件名称、版本、开发完成日期和发表事实；",
        "- 著作权人、开发方式、权利取得方式和权属文件；",
        "- 代码范围、第三方依赖、截图隐私和最终排版；",
        "- 当前受理渠道的最新要求。",
    ])
    if report.blockers:
        lines.extend(["", "## 阻断项", ""] + [f"- {x}" for x in report.blockers])
    path.write_text("\n".join(lines) + "\n", encoding="utf-8")
    return path


def docx_contains(path: Path, values: Sequence[str]) -> bool:
    if not path.exists():
        return False
    try:
        with zipfile.ZipFile(path) as zf:
            text = "\n".join(
                zf.read(name).decode("utf-8", errors="ignore")
                for name in zf.namelist()
                if name.endswith(".xml")
            )
        return all(value in text for value in values if value)
    except Exception:
        return False


def comprehensive_validate(repo: Path, config_path: Path, manual_path: Path, evidence_path: Path, out: Path) -> Report:
    ensure_out(out)
    config = load_toml(config_path)
    evidence = load_json(evidence_path)
    report = Report()
    report.merge(validate_required_facts(config))
    git = git_info(repo, [out])
    report.merge(validate_git(config, git))
    bundle, source_report = collect_source_bundle(repo, config, git)
    report.merge(source_report)
    report.merge(validate_evidence(repo, evidence, str(cfg(config, "software", "manual_type", ""))))

    manual_text = manual_path.read_text(encoding="utf-8") if manual_path.exists() else ""
    if not manual_path.exists():
        report.blockers.append(f"说明书 Markdown 不存在：{manual_path}")
    else:
        placeholders = sorted(set(m.group(0) for m in PLACEHOLDER_RE.finditer(manual_text)))
        if placeholders:
            report.blockers.append("说明书含占位符：" + "、".join(placeholders[:10]))
        full_name = str(cfg(config, "software", "full_name", ""))
        version = str(cfg(config, "software", "version", ""))
        rendered_text = manual_text.replace("{{SOFTWARE_FULL_NAME}}", full_name).replace("{{VERSION}}", version)
        if full_name and full_name not in rendered_text:
            report.blockers.append("说明书正文未出现完整的软件全称。")
        if version and version not in rendered_text:
            report.blockers.append("说明书正文未出现版本号。")
        for feature in evidence.get("features", []):
            if not isinstance(feature, dict):
                continue
            name = str(feature.get("name", "")).strip()
            status = feature.get("status")
            if status == "verified" and name and name not in rendered_text:
                report.warnings.append(f"已核验功能未在说明书中明确出现：{name}")
            if status == "excluded" and name and name in rendered_text:
                report.blockers.append(f"已排除功能仍出现在说明书中：{name}")

    all_paths = tracked_files(repo, git)
    findings, secret_report = scan_secrets(repo, bundle.source_files, all_paths)
    report.merge(secret_report)
    write_secret_findings(findings, out)

    full_name = str(cfg(config, "software", "full_name", ""))
    version = str(cfg(config, "software", "version", ""))
    source_docx = out / "02_source-code-submission.docx"
    manual_docx = out / "03_software-manual.docx"
    if not source_docx.exists():
        report.blockers.append("缺少源程序提交版 DOCX。")
    elif not docx_contains(source_docx, [full_name, version]):
        report.blockers.append("源程序 DOCX 中未检测到一致的软件名称和版本页眉。")
    if not manual_docx.exists():
        report.blockers.append("缺少软件说明书 DOCX。")
    elif not docx_contains(manual_docx, [full_name, version]):
        report.blockers.append("说明书 DOCX 中未检测到一致的软件名称和版本。")

    source_pdf = out / "02_source-code-submission.pdf"
    if source_pdf.exists():
        pages = pdf_page_count(source_pdf)
        if pages is None:
            report.warnings.append("无法读取源程序 PDF 页数。")
        else:
            if pages != bundle.expected_pages:
                report.blockers.append(f"源程序 PDF 页数为 {pages}，预期 {bundle.expected_pages}。")
            if pages > 60:
                report.blockers.append(f"源程序提交版超过60页：{pages}页。")
    elif bool(cfg(config, "output", "generate_pdf", True)):
        report.warnings.append("未生成源程序 PDF，请人工转换并逐页检查。")

    manual_full_pdf = out / "03_software-manual-full.pdf"
    manual_submission_pdf = out / "03_software-manual-submission.pdf"
    if manual_full_pdf.exists():
        full_pages = pdf_page_count(manual_full_pdf)
        submission_pages = pdf_page_count(manual_submission_pdf) if manual_submission_pdf.exists() else None
        if full_pages and full_pages > 60 and submission_pages != 60:
            report.blockers.append("完整说明书超过60页，但提交版不是60页。")
        if full_pages and full_pages <= 60 and submission_pages and submission_pages != full_pages:
            report.blockers.append("说明书不足60页，但提交版页数与完整版不一致。")
    elif bool(cfg(config, "output", "generate_pdf", True)):
        report.warnings.append("未生成说明书 PDF，请人工转换并逐页检查。")

    required_outputs = [
        "01_software-info-draft.json", "01_software-info-draft.md", "02_source-file-manifest.csv",
        "04_feature-evidence.json", "05_dependency-and-ip-report.md",
    ]
    for name in required_outputs:
        if not (out / name).exists():
            report.blockers.append(f"缺少输出文件：{name}")

    report.info.extend([
        "已检查必填事实、日期和交存模式。",
        "已检查 Git 快照和工作区状态。",
        "已检查源代码范围、分页和敏感信息。",
        "已检查功能证据、说明书占位符及名称版本一致性。",
        "已检查 DOCX/PDF 输出存在性和可读取页数。",
    ])
    write_validation_report(report, config, git, out)
    write_human_checklist(config, report, out)
    write_run_summary(config, git, report, out)
    return report


def build_package_manifest(out: Path, git: GitInfo) -> Path:
    files = []
    for path in sorted(out.iterdir(), key=lambda p: p.name):
        if not path.is_file() or path.name in {"software-copyright-technical-package.zip", "package-manifest.json"}:
            continue
        files.append({
            "path": path.name,
            "sha256": sha256_file(path),
            "bytes": path.stat().st_size,
        })
    manifest = {
        "tool_version": TOOL_VERSION,
        "generated_at": dt.datetime.now(dt.timezone.utc).astimezone().isoformat(timespec="seconds"),
        "git_commit": git.commit,
        "files": files,
    }
    path = out / "package-manifest.json"
    path.write_text(json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")
    return path


def create_zip(out: Path) -> Path:
    zip_path = out / "software-copyright-technical-package.zip"
    with zipfile.ZipFile(zip_path, "w", compression=zipfile.ZIP_DEFLATED) as zf:
        for path in sorted(out.iterdir(), key=lambda p: p.name):
            if not path.is_file() or path == zip_path:
                continue
            zf.write(path, arcname=path.name)
    return zip_path


def cmd_init(args: argparse.Namespace) -> int:
    repo = args.repo.resolve()
    skill_root = Path(__file__).resolve().parent.parent
    target = repo / ".softcopyright"
    screenshots = target / "screenshots"
    screenshots.mkdir(parents=True, exist_ok=True)
    mapping = {
        skill_root / "assets" / "softcopyright.toml.example": target / "softcopyright.toml",
        skill_root / "assets" / "feature-evidence.example.json": target / "feature-evidence.json",
        skill_root / "assets" / "manual-template.md": target / "manual.md",
    }
    for src, dst in mapping.items():
        if dst.exists() and not args.force:
            print(f"保留已有文件：{dst}")
        else:
            shutil.copy2(src, dst)
            print(f"已创建：{dst}")
    gitignore = target / ".gitignore"
    if not gitignore.exists():
        gitignore.write_text("screenshots/private/\n", encoding="utf-8")
    print("下一步：由申请人员人工确认配置、功能证据和说明书正文；本脚本只执行机械整理与校验。")
    return 0


def cmd_inspect(args: argparse.Namespace) -> int:
    _state, report = build_inspection(args.repo.resolve(), args.config.resolve(), args.evidence.resolve(), args.out.resolve(), run_commands=not args.skip_commands)
    config = load_toml(args.config.resolve())
    git = git_info(args.repo.resolve(), [args.out.resolve()])
    write_validation_report(report, config, git, args.out.resolve())
    write_human_checklist(config, report, args.out.resolve())
    write_run_summary(config, git, report, args.out.resolve())
    print(report.status)
    return 2 if report.blockers else 0


def cmd_build_source(args: argparse.Namespace) -> int:
    repo = args.repo.resolve()
    config = load_toml(args.config.resolve())
    out = args.out.resolve()
    ensure_out(out)
    git = git_info(repo, [out])
    bundle, report = collect_source_bundle(repo, config, git)
    write_source_manifest(bundle, out)
    docx_path, build_report = build_source_docx(bundle, config, out)
    report.merge(build_report)
    if docx_path.exists() and bool(cfg(config, "output", "generate_pdf", True)):
        pdf_path = out / "02_source-code-submission.pdf"
        ok, message = convert_docx_to_pdf(docx_path, pdf_path)
        if not ok:
            report.warnings.append(f"源程序 DOCX 转 PDF 失败：{message}")
        else:
            pages = pdf_page_count(pdf_path)
            if pages is not None and pages != bundle.expected_pages:
                report.blockers.append(f"源程序 PDF 页数为 {pages}，预期 {bundle.expected_pages}。")
            else:
                report.info.append(f"已生成源程序 PDF，共 {pages or bundle.expected_pages} 页。")
    config_report = validate_required_facts(config)
    report.merge(config_report)
    write_validation_report(report, config, git, out)
    write_human_checklist(config, report, out)
    write_run_summary(config, git, report, out)
    print(report.status)
    return 2 if report.blockers else 0


def cmd_build_manual(args: argparse.Namespace) -> int:
    repo = args.repo.resolve()
    config = load_toml(args.config.resolve())
    evidence = load_json(args.evidence.resolve())
    out = args.out.resolve()
    ensure_out(out)
    report = validate_required_facts(config)
    report.merge(validate_evidence(repo, evidence, str(cfg(config, "software", "manual_type", ""))))
    docx_path = out / "03_software-manual.docx"
    report.merge(markdown_to_docx(args.manual.resolve(), config, docx_path))
    copy_normalized_evidence(evidence, out)
    if docx_path.exists() and bool(cfg(config, "output", "generate_pdf", True)):
        full_pdf = out / "03_software-manual-full.pdf"
        ok, message = convert_docx_to_pdf(docx_path, full_pdf)
        if not ok:
            report.warnings.append(f"说明书 DOCX 转 PDF 失败：{message}")
        else:
            submission = out / "03_software-manual-submission.pdf"
            count, mode = split_manual_pdf(full_pdf, submission)
            if count is None:
                report.warnings.append(f"说明书 PDF 提交版生成失败：{mode}")
            else:
                report.info.append(f"说明书完整版 {count} 页，提交版模式：{mode}。")
    git = git_info(repo, [out])
    write_validation_report(report, config, git, out)
    write_human_checklist(config, report, out)
    write_run_summary(config, git, report, out)
    print(report.status)
    return 2 if report.blockers else 0


def cmd_validate(args: argparse.Namespace) -> int:
    report = comprehensive_validate(
        args.repo.resolve(), args.config.resolve(), args.manual.resolve(), args.evidence.resolve(), args.out.resolve()
    )
    print(report.status)
    return 2 if report.blockers else 0


def cmd_package(args: argparse.Namespace) -> int:
    repo = args.repo.resolve()
    config_path = args.config.resolve()
    evidence_path = args.evidence.resolve()
    manual_path = args.manual.resolve()
    out = args.out.resolve()
    ensure_out(out)

    _state, inspect_report = build_inspection(repo, config_path, evidence_path, out, run_commands=not args.skip_commands)
    config = load_toml(config_path)
    git = git_info(repo, [out])
    bundle, source_report = collect_source_bundle(repo, config, git)
    inspect_report.merge(source_report)
    docx_path, source_build_report = build_source_docx(bundle, config, out)
    inspect_report.merge(source_build_report)
    if docx_path.exists() and bool(cfg(config, "output", "generate_pdf", True)):
        ok, message = convert_docx_to_pdf(docx_path, out / "02_source-code-submission.pdf")
        if not ok:
            inspect_report.warnings.append(f"源程序 DOCX 转 PDF 失败：{message}")

    manual_docx = out / "03_software-manual.docx"
    inspect_report.merge(markdown_to_docx(manual_path, config, manual_docx))
    if manual_docx.exists() and bool(cfg(config, "output", "generate_pdf", True)):
        full_pdf = out / "03_software-manual-full.pdf"
        ok, message = convert_docx_to_pdf(manual_docx, full_pdf)
        if not ok:
            inspect_report.warnings.append(f"说明书 DOCX 转 PDF 失败：{message}")
        else:
            _count, split_message = split_manual_pdf(full_pdf, out / "03_software-manual-submission.pdf")
            if _count is None:
                inspect_report.warnings.append(f"说明书提交版生成失败：{split_message}")

    validation_report = comprehensive_validate(repo, config_path, manual_path, evidence_path, out)
    validation_report.merge(inspect_report)
    # Deduplicate while preserving order.
    validation_report.blockers = list(dict.fromkeys(validation_report.blockers))
    validation_report.warnings = list(dict.fromkeys(validation_report.warnings))
    validation_report.info = list(dict.fromkeys(validation_report.info))
    write_validation_report(validation_report, config, git, out)
    write_human_checklist(config, validation_report, out)
    write_run_summary(config, git, validation_report, out)
    build_package_manifest(out, git)
    if validation_report.blockers:
        print("BLOCKED")
        return 2
    zip_path = create_zip(out)
    print(f"READY_FOR_HUMAN_REVIEW\n{zip_path}")
    return 0


def add_common(parser: argparse.ArgumentParser, *, manual: bool = False, evidence: bool = False) -> None:
    parser.add_argument("--repo", type=Path, default=Path("."), help="项目根目录")
    parser.add_argument("--config", type=Path, default=Path(".softcopyright/softcopyright.toml"), help="配置文件")
    if manual:
        parser.add_argument("--manual", type=Path, default=Path(".softcopyright/manual.md"), help="说明书 Markdown")
    if evidence:
        parser.add_argument("--evidence", type=Path, default=Path(".softcopyright/feature-evidence.json"), help="功能证据 JSON")
    parser.add_argument("--out", type=Path, default=Path("artifacts/software-copyright"), help="输出目录")


def build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--version", action="version", version=TOOL_VERSION)
    sub = parser.add_subparsers(dest="command", required=True)

    p_init = sub.add_parser("init", help="初始化项目输入模板")
    p_init.add_argument("--repo", type=Path, default=Path("."))
    p_init.add_argument("--force", action="store_true", help="覆盖已有模板")
    p_init.set_defaults(func=cmd_init)

    p_inspect = sub.add_parser("inspect", help="扫描仓库、证据、依赖和敏感信息")
    add_common(p_inspect, evidence=True)
    p_inspect.add_argument("--skip-commands", action="store_true", help="不运行构建/测试/截图命令")
    p_inspect.set_defaults(func=cmd_inspect)

    p_source = sub.add_parser("build-source", help="生成源程序鉴别材料")
    add_common(p_source)
    p_source.set_defaults(func=cmd_build_source)

    p_manual = sub.add_parser("build-manual", help="生成软件说明书")
    add_common(p_manual, manual=True, evidence=True)
    p_manual.set_defaults(func=cmd_build_manual)

    p_validate = sub.add_parser("validate", help="执行完整一致性校验")
    add_common(p_validate, manual=True, evidence=True)
    p_validate.set_defaults(func=cmd_validate)

    p_package = sub.add_parser("package", help="生成、校验并打包全部技术材料")
    add_common(p_package, manual=True, evidence=True)
    p_package.add_argument("--skip-commands", action="store_true", help="不运行构建/测试/截图命令")
    p_package.set_defaults(func=cmd_package)
    return parser


def main() -> int:
    parser = build_parser()
    args = parser.parse_args()
    # Resolve config/manual/evidence/out relative to repo for predictable invocation.
    if hasattr(args, "repo"):
        repo = args.repo.resolve()
        for attr in ("config", "manual", "evidence", "out"):
            if hasattr(args, attr):
                value = getattr(args, attr)
                if isinstance(value, Path) and not value.is_absolute():
                    setattr(args, attr, repo / value)
    try:
        return int(args.func(args))
    except SoftCopyrightError as exc:
        eprint(f"ERROR: {exc}")
        return 2
    except KeyboardInterrupt:
        eprint("已取消。")
        return 130
    except Exception as exc:
        eprint(f"UNEXPECTED ERROR: {type(exc).__name__}: {exc}")
        return 3


if __name__ == "__main__":
    raise SystemExit(main())
