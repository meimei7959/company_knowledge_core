#!/usr/bin/env python3
"""Build a one-to-one China software-copyright registration work package.

Agent-safe `collect` only gathers and validates. `finalize-human` is reserved for a
human applicant after confirming the current official form and actual AI usage.
"""
from __future__ import annotations

import argparse
import csv
import datetime as dt
import hashlib
import html
import json
import re
import shutil
import subprocess
import sys
import zipfile
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Iterable

try:
    import tomllib
except ModuleNotFoundError:  # Python < 3.11
    import tomli as tomllib

SCRIPT_DIR = Path(__file__).resolve().parent
SKILL_ROOT = SCRIPT_DIR.parent
sys.path.insert(0, str(SCRIPT_DIR))
import softcopyright as engine  # noqa: E402

VERSION = "2.1.0"
PUBLIC_DIRS = [
    "00_申报总览", "01_正式上报材料", "02_内部复核佐证_通常不上报",
    "03_待人工补齐", "04_合规记录",
]
WORK_ZIP = "软件著作权_工作资料包.zip"
FORMAL_ZIP = "软件著作权_正式上报包.zip"
DEV_LABELS = {
    "independent": "独立开发", "cooperative": "合作开发",
    "commissioned": "委托开发", "task": "下达任务开发",
}
RIGHT_LABELS = {
    "original": "原始取得", "inherited": "继承取得",
    "assigned": "受让取得", "undertaken": "承受取得",
}
APPLICANT_LABELS = {
    "enterprise": "企业法人或其他组织", "individual": "自然人", "other": "其他主体",
}


@dataclass
class Gate:
    passed: bool = False
    conflict: bool = False
    unresolved: bool = True
    blockers: list[str] = field(default_factory=list)
    warnings: list[str] = field(default_factory=list)
    values: dict[str, Any] = field(default_factory=dict)


@dataclass
class Item:
    item_id: str
    name: str
    applicability: str
    submit: str
    target: str
    source: str
    status: str
    action: str
    files: list[str] = field(default_factory=list)


class PackError(RuntimeError):
    pass


def now() -> str:
    return dt.datetime.now().astimezone().isoformat(timespec="seconds")


def write_text(path: Path, text: str) -> Path:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(text.rstrip() + "\n", encoding="utf-8")
    return path


def load_toml(path: Path) -> dict[str, Any]:
    if not path.is_file():
        raise PackError(f"缺少 TOML：{path}")
    with path.open("rb") as f:
        return tomllib.load(f)


def load_json(path: Path) -> dict[str, Any]:
    if not path.is_file():
        raise PackError(f"缺少 JSON：{path}")
    try:
        return json.loads(path.read_text(encoding="utf-8"))
    except json.JSONDecodeError as exc:
        raise PackError(f"JSON 格式错误：{path}: {exc}") from exc


def as_list(value: Any) -> list[str]:
    if isinstance(value, list):
        return [str(x).strip() for x in value if str(x).strip()]
    if value is None or not str(value).strip():
        return []
    return [str(value).strip()]


def resolve(repo: Path, value: Any) -> Path | None:
    text = str(value or "").strip()
    if not text:
        return None
    p = Path(text)
    return p.resolve() if p.is_absolute() else (repo / p).resolve()


def sha256(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as f:
        for chunk in iter(lambda: f.read(1024 * 1024), b""):
            h.update(chunk)
    return h.hexdigest()


def clean(out: Path) -> None:
    out.mkdir(parents=True, exist_ok=True)
    for name in PUBLIC_DIRS + [WORK_ZIP, FORMAL_ZIP, ".engine"]:
        p = out / name
        if p.is_dir():
            shutil.rmtree(p)
        elif p.exists():
            p.unlink()
    (out / ".engine").mkdir(parents=True, exist_ok=True)


def init_project(repo: Path, force: bool) -> None:
    base = repo / ".softcopyright"
    (base / "screenshots").mkdir(parents=True, exist_ok=True)
    (base / "company-materials").mkdir(parents=True, exist_ok=True)
    mapping = {
        SKILL_ROOT / "assets" / "softcopyright.toml.example": base / "softcopyright.toml",
        SKILL_ROOT / "assets" / "feature-evidence.example.json": base / "feature-evidence.json",
        SKILL_ROOT / "assets" / "manual-template.md": base / "manual.md",
        SKILL_ROOT / "assets" / "authorship-attestation.toml.example": base / "authorship-attestation.toml",
        SKILL_ROOT / "assets" / "company-materials.example.json": base / "company-materials.json",
        SKILL_ROOT / "assets" / "company-materials" / "README.md": base / "company-materials" / "README.md",
    }
    for src, dst in mapping.items():
        if dst.exists() and not force:
            print(f"保留已有文件：{dst}")
        else:
            dst.parent.mkdir(parents=True, exist_ok=True)
            shutil.copy2(src, dst)
            print(f"已创建：{dst}")
    write_text(base / ".gitignore", "screenshots/private/\ncompany-materials/private/\n*身份证*\n*handler-id*\n")
    print("下一步：人工填写配置、说明书、功能证据、公司材料清单和 AI 使用确认。")


def run_engine(repo: Path, out: Path, command: str, config: Path, manual: Path, evidence: Path, skip: bool) -> tuple[int, str]:
    args = [sys.executable, str(SCRIPT_DIR / "softcopyright.py"), command,
            "--repo", str(repo), "--config", str(config), "--out", str(out / ".engine")]
    if command == "inspect":
        args += ["--evidence", str(evidence)]
    else:
        args += ["--manual", str(manual), "--evidence", str(evidence)]
    if skip:
        args.append("--skip-commands")
    try:
        result = subprocess.run(args, cwd=str(repo), stdout=subprocess.PIPE, stderr=subprocess.STDOUT,
                                text=True, check=False, timeout=1800)
    except subprocess.TimeoutExpired as exc:
        output = (exc.stdout or "") + "\n内部引擎超时。"
        write_text(out / ".engine" / "wrapper-engine.log", output)
        return 124, output
    write_text(out / ".engine" / "wrapper-engine.log", result.stdout)
    return result.returncode, result.stdout


def parse_engine(out: Path) -> tuple[str, list[str], list[str]]:
    p = out / ".engine" / "06_consistency-validation.md"
    if not p.is_file():
        return "BLOCKED", ["内部引擎未生成一致性报告。"], []
    text = p.read_text(encoding="utf-8", errors="replace")
    m = re.search(r"状态：\*\*(.*?)\*\*", text)
    status = m.group(1).strip() if m else "BLOCKED"
    def section(name: str) -> list[str]:
        mm = re.search(rf"## {re.escape(name)}\s*\n(.*?)(?=\n## |\Z)", text, re.S)
        if not mm:
            return []
        return [x.strip()[2:] for x in mm.group(1).splitlines()
                if x.strip().startswith("- ") and x.strip() != "- 无。"]
    return status, section("阻断项"), section("警告")


def ai_gate(data: dict[str, Any]) -> Gate:
    a = data.get("attestation", {}) if isinstance(data, dict) else {}
    g = Gate(values=dict(a))
    if not bool(a.get("reviewed_current_official_form", False)):
        g.blockers.append("尚未由申请负责人核对当前官方申请表及其 AI 使用声明。")
    for key, label in (
        ("code_ai_used", "源代码编写"),
        ("document_ai_used", "说明书/文档撰写"),
        ("registration_material_ai_used", "登记申请材料生成"),
    ):
        v = str(a.get(key, "unknown")).strip().lower()
        if v not in {"yes", "no", "unknown"}:
            g.blockers.append(f"{key} 必须为 yes、no 或 unknown。")
            v = "unknown"
        g.values[key] = v
        if v == "yes":
            g.conflict = True
            g.blockers.append(f"已确认 AI 参与{label}，与当前公开申请表声明存在冲突。")
        elif v == "unknown":
            g.warnings.append(f"尚未确认 AI 是否参与{label}。")
    for key, label in (
        ("manual_human_authored", "说明书正文由人独立撰写"),
        ("feature_evidence_human_authored", "功能证据叙述由人独立撰写"),
        ("application_text_human_authored", "申请表文字由人独立撰写"),
    ):
        if not bool(a.get(key, False)):
            g.blockers.append(f"未确认：{label}。")
    if not str(a.get("confirmed_by", "")).strip():
        g.blockers.append("AI 使用确认缺少 confirmed_by。")
    date_text = str(a.get("confirmed_on", "")).strip()
    if not date_text:
        g.blockers.append("AI 使用确认缺少 confirmed_on。")
    else:
        try:
            if dt.date.fromisoformat(date_text) > dt.date.today():
                g.blockers.append("AI 使用确认日期在未来。")
        except ValueError:
            g.blockers.append("confirmed_on 必须为 YYYY-MM-DD。")
    vals = [g.values.get(k, "unknown") for k in ("code_ai_used", "document_ai_used", "registration_material_ai_used")]
    human_flags = [bool(a.get(k, False)) for k in ("manual_human_authored", "feature_evidence_human_authored", "application_text_human_authored")]
    g.unresolved = (
        any(v != "no" for v in vals)
        or not bool(a.get("reviewed_current_official_form", False))
        or not all(human_flags)
        or not str(a.get("confirmed_by", "")).strip()
        or not str(a.get("confirmed_on", "")).strip()
    )
    g.passed = not g.blockers and all(v == "no" for v in vals)
    return g


def no_space_len(text: str) -> int:
    return len(re.sub(r"\s+", "", text or ""))


def validate_human_inputs(config: dict[str, Any], evidence: dict[str, Any], manual: Path) -> tuple[list[str], list[str]]:
    b: list[str] = []
    w: list[str] = []
    app = config.get("application", {})
    for key, label in (
        ("applicant_type", "申请人类型"), ("development_mode", "开发方式"),
        ("rights_acquisition", "权利取得方式"), ("software_category", "软件分类"),
        ("rights_scope", "权利范围"), ("main_function_description", "主要功能描述"),
        ("human_author_name", "申请文字人工撰写人"), ("human_confirmed_on", "人工确认日期"),
    ):
        if not str(app.get(key, "")).strip():
            b.append(f"缺少人工确认字段 application.{key}（{label}）。")
    if str(app.get("applicant_type", "")) not in APPLICANT_LABELS:
        b.append("application.applicant_type 枚举无效。")
    if str(app.get("development_mode", "")) not in DEV_LABELS:
        b.append("application.development_mode 枚举无效。")
    if str(app.get("rights_acquisition", "")) not in RIGHT_LABELS:
        b.append("application.rights_acquisition 枚举无效。")
    n = no_space_len(str(app.get("main_function_description", "")))
    try:
        min_chars = int(app.get("main_function_min_chars", 500))
        max_chars = int(app.get("main_function_max_chars", 1300))
    except (TypeError, ValueError):
        min_chars, max_chars = 500, 1300
        b.append("application.main_function_min_chars/max_chars 必须为整数。")
    if min_chars < 1 or max_chars < min_chars:
        b.append("主要功能描述字数范围配置无效。")
    elif not min_chars <= n <= max_chars:
        b.append(f"主要功能描述为 {n} 字，未满足人工按当前表单确认的 {min_chars} 至 {max_chars} 字范围。")
    sw = config.get("software", {})
    for key, label in (
        ("copyright_owner", "著作权人"),
        ("hardware_environment", "硬件环境"),
        ("software_environment", "软件环境"),
        ("technical_characteristics", "技术特点"),
    ):
        if not str(sw.get(key, "")).strip():
            b.append(f"缺少人工确认字段 software.{key}（{label}）。")
    confirmed_on = str(app.get("human_confirmed_on", "")).strip()
    if confirmed_on:
        try:
            if dt.date.fromisoformat(confirmed_on) > dt.date.today():
                b.append("application.human_confirmed_on 在未来。")
        except ValueError:
            b.append("application.human_confirmed_on 必须为 YYYY-MM-DD。")
    if not manual.is_file():
        b.append(f"人工原创说明书不存在：{manual}")
    else:
        text = manual.read_text(encoding="utf-8", errors="replace")
        markers = [x for x in ["TODO", "TBD", "待补充", "占位", "{{SOFTWARE", "HUMAN_ONLY"] if x.lower() in text.lower()]
        if markers:
            b.append("说明书仍含模板或占位标记：" + "、".join(markers))
        if no_space_len(text) < 500:
            w.append("说明书正文较短，请人工确认是否完整描述当前软件。")
    feats = evidence.get("features", []) if isinstance(evidence, dict) else []
    if not isinstance(feats, list) or not feats:
        b.append("功能证据列表为空。")
    verified = 0
    for i, f in enumerate(feats if isinstance(feats, list) else [], 1):
        if not isinstance(f, dict):
            b.append(f"第 {i} 项功能不是对象。")
            continue
        name = str(f.get("name", "")).strip() or f"第{i}项功能"
        status = str(f.get("status", "")).strip()
        if status == "excluded":
            continue
        if status != "verified":
            b.append(f"功能未完成核验：{name}（{status or '空'}）。")
            continue
        verified += 1
        if not str(f.get("summary", "")).strip(): b.append(f"功能缺少人工说明：{name}")
        if not as_list(f.get("code_paths")): b.append(f"功能缺少代码路径：{name}")
        if not as_list(f.get("operation_steps")): b.append(f"功能缺少人工操作步骤：{name}")
    if verified == 0:
        b.append("没有 verified 功能。")
    return b, w


def copy_file(src: Path, dst: Path) -> Path:
    dst.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(src, dst)
    return dst


def company_files(repo: Path, manifest: dict[str, Any], formal: Path) -> tuple[dict[str, list[Path]], list[str]]:
    copied = {k: [] for k in ["form", "identity", "contract", "permission", "succession", "agent", "translation", "other"]}
    warnings: list[str] = []
    def one(key: str, cat: str, folder: str, stem: str) -> None:
        src = resolve(repo, manifest.get(key, ""))
        if src is None: return
        if not src.is_file():
            warnings.append(f"公司材料路径不存在：{key} -> {src}")
            return
        copied[cat].append(copy_file(src, formal / folder / f"{stem}{src.suffix.lower()}"))
    def many(key: str, cat: str, folder: str) -> None:
        for i, val in enumerate(as_list(manifest.get(key, [])), 1):
            src = resolve(repo, val)
            if src is None or not src.is_file():
                warnings.append(f"公司材料路径不存在：{key}[{i}] -> {val}")
                continue
            copied[cat].append(copy_file(src, formal / folder / f"{i:02d}_{src.name}"))
    one("official_application_form", "form", "01_软件著作权登记申请表", "01_正式申请表_签章件")
    if str(manifest.get("applicant_type", "enterprise")) == "enterprise":
        one("business_license", "identity", "03_申请人身份证明", "03-1_营业执照副本复印件")
    else:
        one("identity_document", "identity", "03_申请人身份证明", "03-1_申请人身份证明")
    many("development_contract_or_task_documents", "contract", "04_权属及其他证明_按需/01_开发合同或项目任务书")
    many("original_software_permission_documents", "permission", "04_权属及其他证明_按需/02_原软件著作权人许可证明")
    many("succession_or_transfer_documents", "succession", "04_权属及其他证明_按需/03_继承受让承受证明")
    many("agent_authorization_documents", "agent", "04_权属及其他证明_按需/04_代理委托材料")
    many("foreign_document_translations", "translation", "04_权属及其他证明_按需/05_外文材料中文译本")
    many("other_documents", "other", "04_权属及其他证明_按需/06_其他按需材料")
    return copied, warnings


def application_handoff(config: dict[str, Any], metadata: dict[str, Any], folder: Path) -> list[str]:
    sw, app = config.get("software", {}), config.get("application", {})
    tech = metadata.get("technical_draft", {}) if isinstance(metadata, dict) else {}
    desc = str(app.get("main_function_description", ""))
    n = no_space_len(desc)
    try:
        min_chars = int(app.get("main_function_min_chars", 500))
        max_chars = int(app.get("main_function_max_chars", 1300))
    except (TypeError, ValueError):
        min_chars, max_chars = 500, 1300
    data = {
        "notice": "工作文件，不是正式申请表；所有叙述字段均须由人独立撰写并确认。",
        "software_full_name": sw.get("full_name", ""), "software_short_name": sw.get("short_name", ""),
        "version": sw.get("version", ""), "copyright_owner": sw.get("copyright_owner", ""),
        "applicant_type": APPLICANT_LABELS.get(str(app.get("applicant_type", "")), ""),
        "software_category": app.get("software_category", ""),
        "development_completion_date": sw.get("completion_date", ""),
        "publication_status": sw.get("publication_status", ""), "first_publication_date": sw.get("first_publication_date", ""),
        "development_method": DEV_LABELS.get(str(app.get("development_mode", "")), ""),
        "rights_acquisition_method": RIGHT_LABELS.get(str(app.get("rights_acquisition", "")), ""),
        "rights_scope": app.get("rights_scope", ""), "platforms": sw.get("platforms", []),
        "app_display_name": sw.get("app_display_name", ""), "package_id": sw.get("package_id", ""),
        "hardware_environment": sw.get("hardware_environment", ""), "software_environment": sw.get("software_environment", ""),
        "programming_languages": tech.get("programming_languages", []), "source_nonempty_lines": tech.get("source_nonempty_lines", 0),
        "technical_characteristics": sw.get("technical_characteristics", ""),
        "main_function_description": desc, "main_function_description_character_count": n,
        "main_function_expected_range": [min_chars, max_chars],
        "human_author_name": app.get("human_author_name", ""), "human_confirmed_on": app.get("human_confirmed_on", ""),
    }
    folder.mkdir(parents=True, exist_ok=True)
    (folder / "01_申请表字段交接单.json").write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")
    lines = ["# 软件著作权登记申请表字段交接单", "", "> 本文件不是正式申请表。正式表必须从当前版权登记系统生成并按要求签章。", "", "| 字段 | 人工确认值 |", "|---|---|"]
    labels = {
        "software_full_name":"软件全称", "software_short_name":"软件简称", "version":"版本号", "applicant_type":"申请人类型",
        "software_category":"软件分类", "copyright_owner":"著作权人", "development_completion_date":"开发完成日期",
        "publication_status":"发表状态", "first_publication_date":"首次发表日期", "development_method":"开发方式",
        "rights_acquisition_method":"权利取得方式", "rights_scope":"权利范围", "app_display_name":"应用显示名称",
        "package_id":"包名/标识", "hardware_environment":"硬件环境", "software_environment":"软件环境",
        "programming_languages":"编程语言", "source_nonempty_lines":"源程序非空行数", "technical_characteristics":"技术特点",
        "human_author_name":"申请文字人工撰写人", "human_confirmed_on":"人工确认日期",
    }
    for key, label in labels.items():
        v = data.get(key, "")
        if isinstance(v, list): v = "、".join(str(x) for x in v)
        lines.append(f"| {label} | {str(v).replace('|','｜')} |")
    lines += ["", f"## 主要功能描述（{n} 字）", "", desc or "（尚未由人填写）", "", "## 官网人工处理", "", "- 使用当前版本申请表；", "- 核对诚信和 AI 使用声明；", "- 填写经办人信息；", "- 完成签名或盖章；", "- 将正式表放入公司材料目录并重新最终化。"]
    write_text(folder / "01_申请表字段交接单.md", "\n".join(lines))
    return [] if min_chars <= n <= max_chars else [f"主要功能描述当前为 {n} 字，人工确认范围为 {min_chars} 至 {max_chars} 字。"]


def set_doc_fonts(doc: Any) -> None:
    from docx.oxml.ns import qn
    from docx.shared import Pt
    styles = doc.styles
    for name in ["Normal", "Title", "Heading 1", "Heading 2"]:
        if name in styles:
            style = styles[name]
            style.font.name = "Arial"
            style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    styles["Normal"].font.size = Pt(10.5)


def add_doc_header(doc: Any, full_name: str, version: str) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    for section in doc.sections:
        p = section.header.paragraphs[0]
        p.text = f"{full_name} {version}"
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer.add_run("第 ")
        fld = OxmlElement("w:fldSimple"); fld.set(qn("w:instr"), "PAGE"); footer._p.append(fld)
        footer.add_run(" 页")


def image_width(path: Path, max_inches: float = 6.0) -> float:
    try:
        from PIL import Image
        with Image.open(path) as im:
            w, h = im.size
        return min(max_inches, max(1.0, (w / max(h, 1)) * 6.2))
    except Exception:
        return max_inches


def build_internal_docs(repo: Path, config: dict[str, Any], evidence: dict[str, Any], out: Path) -> tuple[list[str], list[str]]:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches
    internal = out / "02_内部复核佐证_通常不上报"
    shots = internal / "01_界面截图原图"; shots.mkdir(parents=True, exist_ok=True)
    sw = config.get("software", {}); full = str(sw.get("full_name", "")); version = str(sw.get("version", ""))
    warnings: list[str] = []; blockers: list[str] = []; rows: list[dict[str, Any]] = []
    shot_records: list[tuple[str, Path, str]] = []
    features = evidence.get("features", []) if isinstance(evidence, dict) else []
    if not isinstance(features, list): features = []
    for i, f in enumerate(features, 1):
        if not isinstance(f, dict):
            blockers.append(f"第 {i} 项功能证据不是对象。")
            continue
        name = str(f.get("name", "")).strip() or f"功能{i}"
        code_paths = as_list(f.get("code_paths")); screenshots = as_list(f.get("screenshots")); steps = as_list(f.get("operation_steps"))
        copied: list[str] = []
        for j, ref in enumerate(screenshots, 1):
            src = resolve(repo, ref)
            if src is None or not src.is_file():
                warnings.append(f"功能 {name} 的截图不存在：{ref}")
                continue
            safe = re.sub(r"[^0-9A-Za-z\u4e00-\u9fff._-]+", "_", src.name)
            dst = shots / f"{i:02d}-{j:02d}_{safe}"; copy_file(src, dst)
            copied.append(dst.relative_to(out).as_posix()); shot_records.append((name, dst, str(f.get("summary", ""))))
        rows.append({"index":i, "name":name, "status":str(f.get("status", "")), "summary":str(f.get("summary", "")),
                     "code_paths":code_paths, "screenshots":screenshots, "copied":copied, "steps":steps, "notes":str(f.get("notes", ""))})
    (internal / "04_功能代码截图对应表.json").write_text(json.dumps({"features":rows}, ensure_ascii=False, indent=2), encoding="utf-8")
    with (internal / "04_功能代码截图对应表.csv").open("w", encoding="utf-8-sig", newline="") as f:
        w = csv.writer(f); w.writerow(["序号","功能模块","状态","人工功能说明","代码路径","原始截图","归档截图","人工操作步骤","备注"])
        for r in rows: w.writerow([r["index"],r["name"],r["status"],r["summary"],"\n".join(r["code_paths"]),"\n".join(r["screenshots"]),"\n".join(r["copied"]),"\n".join(r["steps"]),r["notes"]])
    # Screenshot booklet DOCX
    doc = Document(); set_doc_fonts(doc); add_doc_header(doc, full, version)
    t = doc.add_heading(f"{full} {version} 界面截图册", 0); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("内部复核材料，通常不单独上报。全部图片应来自真实运行版本并完成隐私脱敏。")
    if not shot_records: doc.add_paragraph("当前没有可归档的真实运行截图。")
    for i, (name, path, summary) in enumerate(shot_records, 1):
        doc.add_heading(f"{i}. {name}", level=1); doc.add_paragraph(summary or "（未提供人工说明）")
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(path), width=Inches(image_width(path)))
    doc.save(internal / "02_界面截图册.docx")
    # Function-module DOCX
    doc2 = Document(); set_doc_fonts(doc2); add_doc_header(doc2, full, version)
    title = doc2.add_heading(f"{full} {version} 功能模块说明书", 0); title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc2.add_paragraph("内部复核材料，通常不单独上报。内容机械汇编自经人工确认的功能证据。")
    table = doc2.add_table(rows=1, cols=4); table.style = "Table Grid"
    for cell, val in zip(table.rows[0].cells, ["序号","功能模块","状态","人工功能说明"]): cell.text = val
    for r in rows:
        cells = table.add_row().cells
        for c, val in zip(cells, [str(r["index"]), r["name"], r["status"], r["summary"] or "（未提供）"]): c.text = val
    for r in rows:
        doc2.add_heading(f"{r['index'] + 1}. {r['name']}", level=1)
        doc2.add_heading("功能说明", level=2); doc2.add_paragraph(r["summary"] or "（未提供人工说明）")
        doc2.add_heading("对应代码路径", level=2)
        for x in r["code_paths"] or ["（未提供）"]: doc2.add_paragraph(x, style="List Bullet")
        doc2.add_heading("操作步骤", level=2)
        for x in r["steps"] or ["（未提供人工步骤）"]: doc2.add_paragraph(x, style="List Number")
        doc2.add_heading("界面证据", level=2)
        for rel in r["copied"]:
            path = out / rel; p = doc2.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run().add_picture(str(path), width=Inches(image_width(path)))
        if not r["copied"]: doc2.add_paragraph("（无已归档截图）")
        if r["notes"]: doc2.add_heading("备注", level=2); doc2.add_paragraph(r["notes"])
    doc2.save(internal / "03_功能模块说明书.docx")
    # Direct PDF using ReportLab, avoiding office-converter locks.
    try:
        from reportlab.lib import colors
        from reportlab.lib.enums import TA_CENTER
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.lib.units import cm
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.cidfonts import UnicodeCIDFont
        from reportlab.platypus import Image as RI, PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle
        pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle("ct", parent=styles["Title"], fontName="STSong-Light", alignment=TA_CENTER, fontSize=18, leading=24)
        h1 = ParagraphStyle("ch1", parent=styles["Heading1"], fontName="STSong-Light", fontSize=14, leading=20)
        body = ParagraphStyle("cb", parent=styles["BodyText"], fontName="STSong-Light", fontSize=9.5, leading=15)
        def rimg(path: Path, mw: float, mh: float):
            im = RI(str(path)); scale = min(mw/im.imageWidth, mh/im.imageHeight, 1.0); im.drawWidth*=scale; im.drawHeight*=scale; return im
        # Screenshot PDF
        story: list[Any] = [Paragraph(html.escape(f"{full} {version} 界面截图册"), title_style), Paragraph("内部复核材料，通常不单独上报。", body)]
        if not shot_records: story.append(Paragraph("当前没有截图。", body))
        for i,(name,path,summary) in enumerate(shot_records,1):
            story += [Paragraph(html.escape(f"{i}. {name}"), h1), Paragraph(html.escape(summary or "（未提供人工说明）"), body), rimg(path,16*cm,19*cm)]
            if i < len(shot_records): story.append(PageBreak())
        SimpleDocTemplate(str(internal/"02_界面截图册.pdf"), pagesize=A4, rightMargin=1.8*cm,leftMargin=1.8*cm,topMargin=1.7*cm,bottomMargin=1.7*cm).build(story)
        # Function PDF
        story = [Paragraph(html.escape(f"{full} {version} 功能模块说明书"), title_style), Paragraph("内部复核材料，通常不单独上报。", body), Paragraph("1. 功能模块总览", h1)]
        data = [[Paragraph(x, body) for x in ["序号","功能模块","状态","功能说明"]]]
        for r in rows: data.append([Paragraph(str(r["index"]),body),Paragraph(html.escape(r["name"]),body),Paragraph(html.escape(r["status"]),body),Paragraph(html.escape(r["summary"] or "（未提供）"),body)])
        tab=Table(data,colWidths=[1.1*cm,3.3*cm,2.2*cm,10.2*cm],repeatRows=1); tab.setStyle(TableStyle([("GRID",(0,0),(-1,-1),0.4,colors.grey),("BACKGROUND",(0,0),(-1,0),colors.lightgrey),("VALIGN",(0,0),(-1,-1),"TOP")]))
        story += [tab,Spacer(1,10)]
        for pos,r in enumerate(rows,1):
            story += [Paragraph(html.escape(f"{pos+1}. {r['name']}"),h1),Paragraph("功能说明",h1),Paragraph(html.escape(r["summary"] or "（未提供）"),body),Paragraph("对应代码路径",h1)]
            for x in r["code_paths"] or ["（未提供）"]: story.append(Paragraph("- "+html.escape(x),body))
            story.append(Paragraph("操作步骤",h1))
            for j,x in enumerate(r["steps"] or ["（未提供）"],1): story.append(Paragraph(f"{j}. {html.escape(x)}",body))
            for rel in r["copied"]: story += [Paragraph("界面证据",h1),rimg(out/rel,15.5*cm,14*cm)]
            if pos < len(rows): story.append(PageBreak())
        SimpleDocTemplate(str(internal/"03_功能模块说明书.pdf"), pagesize=A4, rightMargin=1.8*cm,leftMargin=1.8*cm,topMargin=1.7*cm,bottomMargin=1.7*cm).build(story)
    except Exception as exc:
        warnings.append(f"内部 PDF 生成失败：{exc}")
    return blockers, warnings


def copy_engine_internal(engine_out: Path, internal: Path, finalize: bool) -> None:
    mapping = {
        "00_run-summary.md":"00_内部引擎运行摘要.md", "01_software-info-draft.json":"06_软件基本信息草稿.json",
        "01_software-info-draft.md":"06_软件基本信息草稿.md", "02_source-file-manifest.csv":"05_源码文件清单.csv",
        "04_feature-evidence.json":"04_功能证据原始记录.json", "05_dependency-and-ip-report.md":"07_第三方依赖与IP报告.md",
        "06_consistency-validation.md":"08_一致性校验报告.md", "07_human-review-checklist.md":"09_人工复核清单.md",
    }
    internal.mkdir(parents=True, exist_ok=True)
    for src,dst in mapping.items():
        if (engine_out/src).is_file(): copy_file(engine_out/src, internal/dst)
    if finalize and (engine_out/"03_software-manual-full.pdf").is_file(): copy_file(engine_out/"03_software-manual-full.pdf", internal/"11_说明书完整版.pdf")
    logs = engine_out/".work"/"logs"
    if logs.is_dir(): shutil.copytree(logs, internal/"10_构建测试日志", dirs_exist_ok=True)
    else: write_text(internal/"10_构建测试日志"/"README_未生成日志.md", "未配置构建/测试命令，或本次运行跳过命令。")


def copy_formal_technical(engine_out: Path, config: dict[str, Any], formal: Path) -> tuple[dict[str,list[Path]],list[str]]:
    result = {"source":[],"manual":[]}; warnings=[]; folder=formal/"02_软件鉴别材料"; folder.mkdir(parents=True,exist_ok=True)
    for src,dst in [("02_source-code-submission.docx","02-1_源程序鉴别材料.docx"),("02_source-code-submission.pdf","02-1_源程序鉴别材料.pdf")]:
        if (engine_out/src).is_file(): result["source"].append(copy_file(engine_out/src,folder/dst))
        else: warnings.append(f"缺少源程序输出：{src}")
    label = "用户操作说明书" if str(config.get("software",{}).get("manual_type"))=="user-manual" else "软件设计说明书"
    for src,dst in [("03_software-manual.docx",f"02-2_{label}.docx"),("03_software-manual-submission.pdf",f"02-2_{label}.pdf")]:
        if (engine_out/src).is_file(): result["manual"].append(copy_file(engine_out/src,folder/dst))
        else: warnings.append(f"缺少文档鉴别材料输出：{src}")
    return result,warnings


def conditions(config: dict[str,Any], company: dict[str,Any]) -> dict[str,bool]:
    app=config.get("application",{}); dev=str(app.get("development_mode","")); right=str(app.get("rights_acquisition","")); scenario=str(company.get("ownership_scenario","independent"))
    return {
        "contract": dev in {"cooperative","commissioned","task"} or scenario in {"cooperative","commissioned","task"} or bool(app.get("has_ownership_contract_or_task",False)),
        "permission": bool(app.get("based_on_original_software",False)) or bool(company.get("modified_from_other_software",False)),
        "succession": right not in {"","original"} or bool(company.get("rights_succession",False)),
        "agent": bool(app.get("uses_registration_agent", app.get("uses_agent", False))), "translation": bool(app.get("has_foreign_documents",False)),
    }


def build_items(out: Path, config: dict[str,Any], company: dict[str,Any], c: dict[str,list[Path]], t: dict[str,list[Path]], finalize: bool) -> list[Item]:
    cond=conditions(config,company)
    def rel(paths: Iterable[Path]) -> list[str]: return [p.relative_to(out).as_posix() for p in paths if p.exists()]
    source_ok=finalize and any(p.suffix==".docx" for p in t["source"]) and any(p.suffix==".pdf" for p in t["source"])
    manual_ok=finalize and any(p.suffix==".docx" for p in t["manual"]) and any(p.suffix==".pdf" for p in t["manual"])
    items=[
        Item("F-01","软件著作权登记申请表","必需","是","01_正式上报材料/01_软件著作权登记申请表/01_正式申请表_签章件.*","当前版权登记系统生成＋申请人签章","PRESENT" if c["form"] else "MISSING","在当前系统生成并签章后放入公司材料目录。",rel(c["form"])),
        Item("F-01A","申请表字段交接单","工作文件","否","01_正式上报材料/01_软件著作权登记申请表/01_申请表字段交接单.*","人工确认配置的机械整理","PRESENT","用于人工填官网，不得替代正式表。",rel([out/"01_正式上报材料/01_软件著作权登记申请表/01_申请表字段交接单.md",out/"01_正式上报材料/01_软件著作权登记申请表/01_申请表字段交接单.json"])),
        Item("F-02","源程序鉴别材料","必需","是","01_正式上报材料/02_软件鉴别材料/02-1_源程序鉴别材料.docx/.pdf","真实第一方源码确定性排版","PRESENT" if source_ok else "MISSING","人工最终化并逐页核对。",rel(t["source"])),
        Item("F-03","文档鉴别材料","必需","是","01_正式上报材料/02_软件鉴别材料/02-2_用户操作说明书或软件设计说明书.docx/.pdf","人工原创说明书机械排版","PRESENT" if manual_ok else "MISSING","人工原创并核对真实截图。",rel(t["manual"])),
        Item("F-04","申请人身份证明","必需","是","01_正式上报材料/03_申请人身份证明/","公司/申请人提供","PRESENT" if c["identity"] else "MISSING","补入有效主体身份证明。",rel(c["identity"])),
    ]
    specs=[("F-05","开发合同、合作协议或项目任务书","contract","contract","01_开发合同或项目任务书"),("F-06","原软件著作权人许可证明","permission","permission","02_原软件著作权人许可证明"),("F-07","继承、受让或承受证明","succession","succession","03_继承受让承受证明"),("F-08","代理委托材料","agent","agent","04_代理委托材料"),("F-09","外文证明文件中文译本","translation","translation","05_外文材料中文译本")]
    for fid,name,key,cat,folder in specs:
        applicable=cond[key]; files=c[cat]
        items.append(Item(fid,name,"按需-当前适用" if applicable else "按需-当前不适用","适用时上报",f"01_正式上报材料/04_权属及其他证明_按需/{folder}/","公司/相关权利人提供",("PRESENT" if files else "MISSING") if applicable else "NOT_APPLICABLE","适用时补入并人工确认。" if applicable else "当前配置判断不适用，提交前仍需人工确认。",rel(files)))
    items.append(Item("F-10","其他系统或补正要求材料","按系统提示","按需","01_正式上报材料/04_权属及其他证明_按需/06_其他按需材料/","申请人/公司提供","PRESENT" if c["other"] else "NOT_APPLICABLE","按当前系统或补正通知补充。",rel(c["other"])))
    return items


def applicable_missing(items:list[Item])->tuple[int,int,list[Item]]:
    app=[x for x in items if x.item_id in {"F-01","F-02","F-03","F-04"} or x.applicability=="按需-当前适用"]
    present=[x for x in app if x.status=="PRESENT"]; missing=[x for x in app if x.status!="PRESENT"]
    return len(present),len(app),missing


def write_maps(out:Path,items:list[Item])->None:
    d=out/"00_申报总览"; d.mkdir(parents=True,exist_ok=True)
    with (d/"00_申报材料一一对应表.csv").open("w",encoding="utf-8-sig",newline="") as f:
        w=csv.writer(f); w.writerow(["编号","申报项","适用性","是否上报","目标路径","来源","状态","下一步","实际文件"])
        for x in items: w.writerow([x.item_id,x.name,x.applicability,x.submit,x.target,x.source,x.status,x.action,"\n".join(x.files)])
    lines=["# 软件著作权申报材料一一对应表","","| 编号 | 申报项 | 适用性 | 是否上报 | 状态 | 最终路径 | 下一步 |","|---|---|---|---|---|---|---|"]
    for x in items:
        lines.append(f"| {x.item_id} | {x.name} | {x.applicability} | {x.submit} | **{x.status}** | `{x.target}` | {x.action} |")
        if x.files: lines += ["", "实际文件："+"；".join(f"`{p}`" for p in x.files), ""]
    write_text(d/"00_申报材料一一对应表.md","\n".join(lines))


def determine(finalize:bool,gate:Gate,engine_status:str,engine_blockers:list[str],extra:list[str],items:list[Item])->str:
    _,_,missing=applicable_missing(items)
    if gate.conflict: return "INTERNAL_ONLY_AI_USE"
    if engine_status=="BLOCKED" or engine_blockers or extra: return "BLOCKED"
    if not finalize or not gate.passed or missing: return "READY_FOR_HUMAN_COMPLETION"
    return "SUBMISSION_READY"


def write_state(out:Path,status:str,gate:Gate,engine_status:str,eb:list[str],ew:list[str],xb:list[str],xw:list[str],items:list[Item],git:Any)->None:
    d=out/"00_申报总览"; d.mkdir(parents=True,exist_ok=True); (out/"03_待人工补齐").mkdir(parents=True,exist_ok=True); (out/"04_合规记录").mkdir(parents=True,exist_ok=True)
    p,a,missing=applicable_missing(items)
    state={"tool_version":VERSION,"generated_at":now(),"overall_status":status,"engine_status":engine_status,
           "ai_gate":{"passed":gate.passed,"conflict":gate.conflict,"unresolved":gate.unresolved,"values":gate.values,"blockers":gate.blockers,"warnings":gate.warnings},
           "git":{"commit":git.commit,"branch":git.branch,"dirty":git.dirty},"formal_materials":{"present":p,"applicable":a,"missing":[x.item_id for x in missing]},
           "blockers":list(dict.fromkeys(eb+xb)),"warnings":list(dict.fromkeys(ew+xw+gate.warnings))}
    (d/"00_材料状态.json").write_text(json.dumps(state,ensure_ascii=False,indent=2),encoding="utf-8")
    lines=["# 待人工或公司补齐的申报材料","",f"当前正式申报项：已齐 {p}/{a}。",""]
    for x in missing: lines += [f"## {x.item_id} {x.name}","",f"- 目标路径：`{x.target}`",f"- 来源：{x.source}",f"- 下一步：{x.action}",""]
    if not missing: lines.append("没有缺失的当前适用正式申报项。")
    if gate.blockers or gate.warnings: lines += ["","## AI 合规门",""]+[f"- {x}" for x in gate.blockers+gate.warnings]
    write_text(d/"00_缺件清单.md","\n".join(lines)); write_text(out/"03_待人工补齐/00_请补入以下材料.md","\n".join(lines))
    write_text(d/"00_提交顺序与操作说明.md",f"""# 提交顺序与操作说明

1. 查看一一对应表和缺件清单。
2. 使用字段交接单在当前版权登记系统人工填写。
3. 下载正式申请表，完成当前要求的声明、经办人信息和签章。
4. 补入营业执照/身份证明和适用权属证明。
5. 由申请负责人运行 `finalize-human`，直至状态为 `SUBMISSION_READY`。
6. 按系统槽位逐项上传正式包中的文件，不要把内部佐证区整体上传。
7. 提交前逐页核对软件全称、版本、著作权人、页眉和截图。

当前状态：**{status}**
""")
    gate_lines=["# AI 使用与作者声明检查","","> 内部记录，不替代官网声明。","",f"- 结果：**{'通过' if gate.passed else '未通过/未完成'}**",f"- 已确认冲突：{gate.conflict}",f"- 仍有未知项：{gate.unresolved}","","## 阻断项",""]+[f"- {x}" for x in gate.blockers or ["无。"]]+["","## 警告",""]+[f"- {x}" for x in gate.warnings or ["无。"]]
    write_text(out/"04_合规记录/01_AI使用与作者声明检查.md","\n".join(gate_lines))
    write_text(out/"04_合规记录/00_敏感材料存储提示.md", """# 敏感材料存储提示

工作资料包和正式上报包可能包含营业执照、身份证明、经办人信息、签章、合同及其他敏感材料。ZIP 默认未加密。请存放在受控目录，不上传公开仓库，不通过无保护渠道传播，并按公司数据管理制度设置访问权限和留存期限。
""")
    (out/"04_合规记录/02_Git快照信息.json").write_text(json.dumps({"commit":git.commit,"branch":git.branch,"dirty":git.dirty,"status_lines":git.status_lines,"generated_at":now()},ensure_ascii=False,indent=2),encoding="utf-8")


def hash_manifest(out:Path)->None:
    records=[]
    for rootname in PUBLIC_DIRS:
        root=out/rootname
        if root.exists():
            for p in sorted(root.rglob("*")):
                if p.is_file(): records.append({"path":p.relative_to(out).as_posix(),"sha256":sha256(p),"bytes":p.stat().st_size})
    (out/"04_合规记录/03_文件哈希清单.json").write_text(json.dumps({"tool_version":VERSION,"generated_at":now(),"files":records},ensure_ascii=False,indent=2),encoding="utf-8")


def work_zip(out:Path)->Path:
    target=out/WORK_ZIP
    with zipfile.ZipFile(target,"w",zipfile.ZIP_DEFLATED) as z:
        for rootname in PUBLIC_DIRS:
            root=out/rootname
            if root.exists():
                for p in sorted(root.rglob("*")):
                    if p.is_file(): z.write(p,p.relative_to(out).as_posix())
    return target


def formal_zip(out:Path,items:list[Item])->Path:
    root=out/"01_正式上报材料"; target=out/FORMAL_ZIP
    payload=[]
    for p in sorted(root.rglob("*")):
        if not p.is_file(): continue
        if p.name.startswith("README_") or "申请表字段交接单" in p.name: continue
        payload.append(p)
    hashes=[{"path":p.relative_to(out).as_posix(),"sha256":sha256(p),"bytes":p.stat().st_size} for p in payload]
    present,total,_=applicable_missing(items)
    checklist=["# 正式上报材料清单","",f"当前适用正式申报项已齐：{present}/{total}","","> 实际提交时通常应按系统材料槽位逐项上传。",""]
    for x in items:
        if x.status=="PRESENT" and x.submit in {"是","适用时上报"}: checklist.append(f"- {x.item_id} {x.name}：{'；'.join(x.files)}")
    with zipfile.ZipFile(target,"w",zipfile.ZIP_DEFLATED) as z:
        z.writestr("00_正式上报材料清单.md","\n".join(checklist)+"\n")
        z.writestr("00_正式上报文件哈希.json",json.dumps({"files":hashes},ensure_ascii=False,indent=2))
        for p in payload: z.write(p,p.relative_to(out).as_posix())
    return target


def assemble(args:argparse.Namespace,finalize:bool)->int:
    repo=args.repo.resolve(); out=args.out.resolve(); clean(out)
    config=load_toml(args.config.resolve()); evidence=load_json(args.evidence.resolve()); att=load_toml(args.attestation.resolve()); company=load_json(args.company.resolve()); gate=ai_gate(att)
    rc,_=run_engine(repo,out,"package" if finalize else "inspect",args.config.resolve(),args.manual.resolve(),args.evidence.resolve(),args.skip_commands)
    estatus,eb,ew=parse_engine(out)
    if rc not in {0,2}: eb.append(f"内部引擎异常退出，退出码 {rc}。"); estatus="BLOCKED"
    formal=out/"01_正式上报材料"; internal=out/"02_内部复核佐证_通常不上报"
    for p in [formal,internal,out/"00_申报总览",out/"03_待人工补齐",out/"04_合规记录"]: p.mkdir(parents=True,exist_ok=True)
    write_text(formal/"01_软件著作权登记申请表/README_正式表须由官网生成.md","# 正式申请表\n\n正式表必须从当前版权登记系统生成并由申请人按当前要求签章。Skill 不生成官方表格、签字、公章或身份证号码。")
    write_text(formal/"03_申请人身份证明/README_公司提供.md","# 申请人身份证明\n\n由申请人/公司提供有效主体身份证明，Skill 只原样复制。")
    write_text(formal/"04_权属及其他证明_按需/README_按权属情形提供.md","# 权属及其他证明\n\n仅在对应情形下提供，Skill 不起草或修改合同和证明。")
    metadata={}
    mp=out/".engine/01_software-info-draft.json"
    if mp.is_file():
        try: metadata=json.loads(mp.read_text(encoding="utf-8"))
        except Exception: pass
    handoff_w=application_handoff(config,metadata,formal/"01_软件著作权登记申请表")
    c,cw=company_files(repo,company,formal)
    t={"source":[],"manual":[]}; tw=[]
    if finalize and gate.passed: t,tw=copy_formal_technical(out/".engine",config,formal)
    elif finalize: tw.append("AI 合规门未通过，内部排版结果未复制到正式上报区。")
    copy_engine_internal(out/".engine",internal,finalize)
    ib,iw=build_internal_docs(repo,config,evidence,out)
    hb: list[str]=[]; hw: list[str]=[]
    if finalize: hb,hw=validate_human_inputs(config,evidence,args.manual.resolve())
    xb=[]; xw=list(dict.fromkeys(handoff_w+cw+tw+iw+hw))
    if finalize: xb += ib+hb
    else: xw += ib
    ca=str(config.get("application",{}).get("applicant_type","")); ma=str(company.get("applicant_type",""))
    if ca and ma and ca!=ma: xb.append("softcopyright.toml 与 company-materials.json 的 applicant_type 不一致。")
    items=build_items(out,config,company,c,t,finalize); write_maps(out,items)
    git=engine.git_info(repo,[out]); status=determine(finalize,gate,estatus,eb,xb,items)
    write_state(out,status,gate,estatus,eb,ew,xb,xw,items,git); hash_manifest(out); wz=work_zip(out); fz=None
    if status=="SUBMISSION_READY": fz=formal_zip(out,items)
    p,a,missing=applicable_missing(items)
    print(f"状态：{status}\n软件：{config.get('software',{}).get('full_name','')} {config.get('software',{}).get('version','')}\n代码快照：{git.commit}\n正式申报项：{p}/{a}\nAI合规门：{'通过' if gate.passed else '未通过/未完成'}\n工作资料包：{wz}\n正式上报包：{fz or '未生成'}")
    if missing: print("尚需补齐："+"、".join(f"{x.item_id} {x.name}" for x in missing))
    if eb or xb: print("阻断项："+"；".join(list(dict.fromkeys(eb+xb))))
    return 0 if status in {"SUBMISSION_READY","READY_FOR_HUMAN_COMPLETION","INTERNAL_ONLY_AI_USE"} else 2


def parser()->argparse.ArgumentParser:
    p=argparse.ArgumentParser(description=__doc__); p.add_argument("--version",action="version",version=VERSION); sub=p.add_subparsers(dest="command",required=True)
    pi=sub.add_parser("init"); pi.add_argument("--repo",type=Path,default=Path(".")); pi.add_argument("--force",action="store_true")
    for name in ["collect","finalize-human"]:
        q=sub.add_parser(name); q.add_argument("--repo",type=Path,default=Path(".")); q.add_argument("--config",type=Path,default=Path(".softcopyright/softcopyright.toml")); q.add_argument("--manual",type=Path,default=Path(".softcopyright/manual.md")); q.add_argument("--evidence",type=Path,default=Path(".softcopyright/feature-evidence.json")); q.add_argument("--attestation",type=Path,default=Path(".softcopyright/authorship-attestation.toml")); q.add_argument("--company",type=Path,default=Path(".softcopyright/company-materials.json")); q.add_argument("--out",type=Path,default=Path("artifacts/software-copyright")); q.add_argument("--skip-commands",action="store_true")
    return p


def main()->int:
    args=parser().parse_args(); repo=args.repo.resolve(); args.repo=repo
    for attr in ["config","manual","evidence","attestation","company","out"]:
        if hasattr(args,attr):
            v=getattr(args,attr)
            if isinstance(v,Path) and not v.is_absolute(): setattr(args,attr,repo/v)
    try:
        if args.command=="init": init_project(repo,args.force); return 0
        return assemble(args,args.command=="finalize-human")
    except PackError as exc: print(f"ERROR: {exc}",file=sys.stderr); return 2
    except KeyboardInterrupt: return 130
    except Exception as exc: print(f"UNEXPECTED ERROR: {type(exc).__name__}: {exc}",file=sys.stderr); return 3


if __name__=="__main__": raise SystemExit(main())
